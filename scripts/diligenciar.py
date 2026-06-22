import os
import sys
import json
import datetime
import argparse
import re
import shutil
import openpyxl
from openpyxl.styles import Alignment, Border, Side, PatternFill, Font
import docx
from docx.shared import Pt
from dotenv import load_dotenv

load_dotenv()

# Importar validador de entorno (modulo sibling en scripts/)
_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
if _THIS_DIR not in sys.path:
    sys.path.insert(0, _THIS_DIR)
import env_validator  # noqa: E402

# Caminos absolutos a los archivos de trabajo
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
WORK_DIR = os.path.join(BASE_DIR, "contexto")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
PLANTILLAS_DIR = os.path.join(WORK_DIR, "plantillas")
EXCEL_PATH = os.path.join(PLANTILLAS_DIR, "bitacoras.xlsx")
TEMPLATE_EXCEL_PATH = os.path.join(PLANTILLAS_DIR, "bitacoras.xlsx")
WORD_PATH = os.path.join(PLANTILLAS_DIR, "actas.docx")
TEMPLATE_WORD_PATH = os.path.join(PLANTILLAS_DIR, "actas.docx")
HISTORICO_PATH = os.path.join(WORK_DIR, "historico_actividades.md")
MEMORY_PATH = os.path.join(WORK_DIR, "memory_descriptions.md")

# Constantes de layout
EXCEL_ACT_START_ROW = 40
EXCEL_ACT_ROW_SPAN = 2
EXCEL_MAX_ACTIVITIES = 7
EXCEL_ACT_END_ROW = EXCEL_ACT_START_ROW + (EXCEL_MAX_ACTIVITIES * EXCEL_ACT_ROW_SPAN) - 1  # 53


def parse_memory_descriptions():
    """
    Lee y parsea el bloque JSON contenido dentro del archivo memory_descriptions.md
    """
    if not os.path.exists(MEMORY_PATH):
        raise FileNotFoundError(f"No se encontró el archivo de memoria en {MEMORY_PATH}")

    with open(MEMORY_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extraer el bloque JSON usando regex
    match = re.search(r"```json\s*(.*?)\s*```", content, re.DOTALL)
    if not match:
        raise ValueError("No se encontró el bloque ```json ... ``` en memory_descriptions.md")

    json_str = match.group(1)
    return json.loads(json_str)


def get_next_undiligenced_bitacora():
    """
    Lee historico_actividades.md e identifica el número de la primera bitácora
    que no esté marcada como [DILIGENCIADA].
    """
    if not os.path.exists(HISTORICO_PATH):
        raise FileNotFoundError(f"No se encontró el histórico en {HISTORICO_PATH}")

    with open(HISTORICO_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        match = re.search(r"##\s*[Bb]itacora\s*numero\s*(\d+)", line)
        if match:
            bitacora_num = int(match.group(1))
            if "[DILIGENCIADA]" not in line:
                return bitacora_num
    return None


def get_all_undiligenced_bitacoras():
    """
    Lee historico_actividades.md y retorna una lista con los números de todas
    las bitácoras que no estén marcadas como [DILIGENCIADA].
    """
    if not os.path.exists(HISTORICO_PATH):
        raise FileNotFoundError(f"No se encontró el histórico en {HISTORICO_PATH}")

    with open(HISTORICO_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    pendientes = []
    for line in lines:
        match = re.search(r"##\s*[Bb]itacora\s*numero\s*(\d+)", line)
        if match:
            bitacora_num = int(match.group(1))
            if "[DILIGENCIADA]" not in line:
                pendientes.append(bitacora_num)
    return pendientes


def mark_bitacora_as_diligenced(num):
    """
    Modifica historico_actividades.md para añadir la etiqueta [DILIGENCIADA]
    al encabezado de la bitácora número 'num'.
    """
    if not os.path.exists(HISTORICO_PATH):
        return

    with open(HISTORICO_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    pattern = rf"(##\s*[Bb]itacora\s*numero\s*{num}\s*-[^\n]+)"

    def replacer(match):
        line = match.group(1)
        if "[DILIGENCIADA]" not in line:
            return f"{line.strip()} [DILIGENCIADA]"
        return line

    new_content, count = re.subn(pattern, replacer, content)
    if count > 0:
        with open(HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(new_content)
        print(f"[Estado] historico_actividades.md actualizado: Bitácora {num} marcada como [DILIGENCIADA].")


def copy_style(src_cell, dest_cell):
    """
    Copia el estilo visual de una celda origen a una celda destino
    """
    if src_cell.has_style:
        dest_cell.font = Font(
            name=src_cell.font.name,
            size=src_cell.font.size,
            bold=src_cell.font.bold,
            italic=src_cell.font.italic,
            color=src_cell.font.color,
            underline=src_cell.font.underline
        )
        dest_cell.border = Border(
            left=src_cell.border.left,
            right=src_cell.border.right,
            top=src_cell.border.top,
            bottom=src_cell.border.bottom
        )
        dest_cell.fill = PatternFill(
            fill_type=src_cell.fill.fill_type,
            start_color=src_cell.fill.start_color,
            end_color=src_cell.fill.end_color
        )
        dest_cell.alignment = Alignment(
            horizontal=src_cell.alignment.horizontal,
            vertical=src_cell.alignment.vertical,
            wrap_text=src_cell.alignment.wrap_text
        )
        dest_cell.number_format = src_cell.number_format


def fill_excel_bitacora(bitacora_num, bitacora_data, execution_date, output_dir=None):
    """
    Rellena la bitácora en Excel usando openpyxl.
    Cada actividad ocupa 2 filas con merge vertical en B, C, D, E, F.
    Máximo 7 actividades (filas 40-53).
    """
    # 1. Definir la carpeta de salida
    if output_dir is None:
        date_str = execution_date.strftime("%Y-%m-%d")
        output_dir = os.path.join(OUTPUT_DIR, f"bitacora{bitacora_num}-{date_str}")
    os.makedirs(output_dir, exist_ok=True)

    # 2. Definir el archivo de salida
    output_filename = f"Bitacora_{bitacora_data['fecha_inicio'].replace('/', '-')}_{bitacora_data['fecha_fin'].replace('/', '-')}.xlsx"
    output_path = os.path.join(output_dir, output_filename)

    # 3. Copiar de la plantilla al archivo final
    shutil.copy2(TEMPLATE_EXCEL_PATH, output_path)

    # 4. Abrir y diligenciar
    wb = openpyxl.load_workbook(output_path)
    sheet = wb['GFPI-F-147-Formato Bitácora']

    # Metadatos del encabezado
    sheet['E17'] = bitacora_num
    sheet['F17'] = f"DESDE {bitacora_data['fecha_inicio']}"
    sheet['G17'] = f"HASTA {bitacora_data['fecha_fin']}"
    sheet['F67'] = execution_date

    print(f"[Excel] Diligenciando {output_filename} para Bitácora {bitacora_num}:")
    print(f"        Período: {bitacora_data['fecha_inicio']} - {bitacora_data['fecha_fin']}")
    print(f"        Carpeta: {output_dir}")

    # Limpiar y rellenar actividades (filas 40 a 53)
    # Descombinar cualquier combinación previa en las filas de actividad
    for r in range(EXCEL_ACT_START_ROW, EXCEL_ACT_END_ROW + 1):
        for col in ['B', 'C', 'D', 'E', 'F', 'G']:
            sheet[f"{col}{r}"].value = None
        for rng in list(sheet.merged_cells.ranges):
            if rng.min_row == r and rng.min_col in (2, 4, 5, 6):
                sheet.unmerge_cells(rng.coord)

    # Copiar estilos base de la celda de referencia B40
    ref_cell_b = sheet['B40']
    ref_cell_d = sheet['D40']
    ref_cell_e = sheet['E40']
    ref_cell_f = sheet['F40']
    ref_cell_g = sheet['G40']

    activities = bitacora_data['actividades']
    max_activities_shown = min(len(activities), EXCEL_MAX_ACTIVITIES)

    if len(activities) > EXCEL_MAX_ACTIVITIES:
        print(f"[Warning] Hay {len(activities)} actividades pero el límite en Excel es {EXCEL_MAX_ACTIVITIES}. "
              f"Se mostrarán solo las primeras {EXCEL_MAX_ACTIVITIES}.")

    for idx in range(max_activities_shown):
        act = activities[idx]
        r = EXCEL_ACT_START_ROW + (idx * EXCEL_ACT_ROW_SPAN)
        r2 = r + 1  # segunda fila del par

        # Escribir la descripción en B{r} y combinar B{r}:B{r2}
        sheet[f"B{r}"] = act['descripcion']
        sheet.merge_cells(start_row=r, start_column=2, end_row=r2, end_column=3)
        copy_style(ref_cell_b, sheet[f"B{r}"])
        sheet[f"B{r}"].alignment = Alignment(wrap_text=True, vertical="center", horizontal="left")
        copy_style(ref_cell_b, sheet[f"C{r}"])

        # Merge vertical para D, E, F (de r a r2)
        sheet[f"D{r}"] = act['fecha_inicio']
        sheet.merge_cells(start_row=r, start_column=4, end_row=r2, end_column=4)
        copy_style(ref_cell_d, sheet[f"D{r}"])

        sheet[f"E{r}"] = act['fecha_fin']
        sheet.merge_cells(start_row=r, start_column=5, end_row=r2, end_column=5)
        copy_style(ref_cell_e, sheet[f"E{r}"])

        sheet[f"F{r}"] = act['evidencia']
        sheet.merge_cells(start_row=r, start_column=6, end_row=r2, end_column=6)
        copy_style(ref_cell_f, sheet[f"F{r}"])
        sheet[f"F{r}"].alignment = Alignment(wrap_text=True, vertical="center", horizontal="left")

        # Celda G vacía con estilo
        sheet[f"G{r}"] = ""
        copy_style(ref_cell_g, sheet[f"G{r}"])

        print(f"        Filas {r}-{r2}: {act['descripcion'][:40]}... | {act['fecha_inicio']} al {act['fecha_fin']}")

    wb.save(output_path)
    print(f"[Excel] Grabado exitosamente: {output_path}")
    return output_path


def apply_calibri_9(element):
    """
    Aplica tipografía Calibri tamaño 9 a un párrafo o celda de tabla de Word.
    Afecta todos los runs existentes del elemento.
    """
    if isinstance(element, docx.text.paragraph.Paragraph):
        for run in element.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
    elif isinstance(element, docx.table._Cell):
        for paragraph in element.paragraphs:
            apply_calibri_9(paragraph)


def process_word_actas(moment, actas_data, execution_date_str, output_dir=None):
    """
    Rellena el momento correspondiente (1, 2 o 3) en el Word.
    """
    if moment not in [1, 2, 3]:
        print(f"[Word] Momento {moment} inválido para actas. Ignorando.")
        return

    # --- Determinar ruta de salida ---
    if output_dir is None:
        output_dir = os.path.join(OUTPUT_DIR, "actas")
    os.makedirs(output_dir, exist_ok=True)
    output_word_path = os.path.join(output_dir, "Actas-Inicio-Medio-Final.docx")

    # --- Copiar desde la plantilla solo si no existe el archivo previo ---
    if not os.path.exists(output_word_path):
        shutil.copy2(TEMPLATE_WORD_PATH, output_word_path)

    print(f"[Word] Diligenciando Acta para Momento {moment} (Fecha ejecución: {execution_date_str}):")
    print(f"        Archivo: {output_word_path}")

    doc = docx.Document(output_word_path)

    if moment == 1:
        # --- MOMENTO 1: PLANEACIÓN ---
        table2 = doc.tables[2]
        fecha_inicio_etapa = os.getenv("FECHA_INICIO_ETAPA", "").strip()
        fecha_fin_etapa = os.getenv("FECHA_FIN_ETAPA", "").strip()
        fecha_afiliacion_arl = os.getenv("FECHA_AFILIACION_ARL", "").strip()
        arl_numero = os.getenv("ARL_NUMERO", "").strip()
        horario = os.getenv("HORARIO_ETAPA", "").strip()

        # R1, C1 — Fecha inicio etapa productiva
        table2.rows[1].cells[1].text = fecha_inicio_etapa
        apply_calibri_9(table2.rows[1].cells[1])

        # R1, C7 — Fecha fin etapa productiva
        table2.rows[1].cells[7].text = fecha_fin_etapa
        apply_calibri_9(table2.rows[1].cells[7])

        # R1, C12 — Fecha afiliación ARL (opcional)
        if fecha_afiliacion_arl:
            table2.rows[1].cells[12].text = fecha_afiliacion_arl
            apply_calibri_9(table2.rows[1].cells[12])

        # R2, C3 — Número de póliza ARL (opcional)
        if arl_numero:
            table2.rows[2].cells[3].text = arl_numero
            apply_calibri_9(table2.rows[2].cells[3])

        # R2, C11 — Horario (opcional)
        if horario:
            table2.rows[2].cells[11].text = horario
            apply_calibri_9(table2.rows[2].cells[11])

        # R6-R9: Contenido inferido para Momento 1 (si existe en JSON)
        momento1_data = actas_data.get('actas', {}).get('momento_1', {})
        if momento1_data:
            # R6: Resultados de aprendizaje
            resultados = momento1_data.get('resultados_aprendizaje', '')
            if resultados:
                table2.rows[6].cells[2].text = resultados
                apply_calibri_9(table2.rows[6].cells[2])
            # R7: Actividades a desarrollar
            actividades = momento1_data.get('actividades_desarrollar', '')
            if actividades:
                table2.rows[7].cells[2].text = actividades
                apply_calibri_9(table2.rows[7].cells[2])
            # R8: Evidencias de aprendizaje
            evidencias = momento1_data.get('evidencias_aprendizaje', '')
            if evidencias:
                table2.rows[8].cells[2].text = evidencias
                apply_calibri_9(table2.rows[8].cells[2])
            # R9: Observaciones adicionales
            obs_adicionales = momento1_data.get('observaciones_adicionales', '')
            if obs_adicionales:
                table2.rows[9].cells[2].text = obs_adicionales
                apply_calibri_9(table2.rows[9].cells[2])
            print("        [Momento 1] Contenido inferido aplicado (R6-R9).")

        # P10 — Fecha de diligenciamiento
        doc.paragraphs[10].text = f"Ciudad: Bogotá D.C.  y fecha de diligenciamiento: {execution_date_str} de forma presencial ___ o virtual   X"
        apply_calibri_9(doc.paragraphs[10])

        print("        [Momento 1] Tabla 2 rellenada (fechas, ARL, horario, contenido inferido R6-R9), párrafo P10 actualizado.")

    elif moment == 2:
        # --- MOMENTO 2: SEGUIMIENTO ---
        # 1. Rellenar marcas X en Tabla 3 (Seguimiento).
        table3 = doc.tables[3]
        for r_idx in list(range(6, 14)) + list(range(17, 22)):
            row = table3.rows[r_idx]
            row.cells[2].text = "X"
            apply_calibri_9(row.cells[2])
            for p in row.cells[2].paragraphs:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 2. Rellenar fecha de momento en Tabla 3, Fila 1, Celda 8
        table3.rows[1].cells[8].text = execution_date_str
        apply_calibri_9(table3.rows[1].cells[8])
        for p in table3.rows[1].cells[8].paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 3. Rellenar SOLO observaciones del aprendiz (P17).
        #    NO rellenar P14 (instructor seguimiento) ni P20 (co-formador)
        doc.paragraphs[17].text = f"  {actas_data['actas']['momento_2']['observaciones_aprendiz']}"
        apply_calibri_9(doc.paragraphs[17])

        # P20: Retroalimentación entre co-formador (inferencia)
        coformador_m2 = actas_data.get('actas', {}).get('momento_2', {}).get('observaciones_coformador', '')
        if coformador_m2:
            doc.paragraphs[20].text = f"  {coformador_m2}"
            apply_calibri_9(doc.paragraphs[20])
            print("        [Momento 2] Retroalimentación co-formador aplicada (P20).")

        # 4. Compromisos de mejora en Tabla 3, columna "Observaciones / Compromisos de mejora"
        compromisos = actas_data['actas']['momento_2'].get('compromisos_mejora', '')
        if compromisos:
            # Escribir en la primera fila técnica (row 6, col 6) y actitudinal (row 17, col 6)
            table3.rows[6].cells[6].text = compromisos
            apply_calibri_9(table3.rows[6].cells[6])
            table3.rows[17].cells[6].text = compromisos
            apply_calibri_9(table3.rows[17].cells[6])

        # 5. Fecha de diligenciamiento en párrafo P22
        doc.paragraphs[22].text = f"Ciudad: Bogotá D.C.  y fecha de diligenciamiento: {execution_date_str} de forma presencial ___ o virtual   X "
        apply_calibri_9(doc.paragraphs[22])

        print("        [Momento 2] Tabla 3 rellenada, observaciones del aprendiz insertadas, retroalimentación co-formador aplicada (P17, P20), "
              "compromisos de mejora agregados, párrafo P22 actualizado.")

    elif moment == 3:
        # --- MOMENTO 3: EVALUACIÓN FINAL ---
        # 1. Rellenar marcas X en Tabla 5 (Evaluación).
        table5 = doc.tables[5]
        for r_idx in list(range(6, 14)) + list(range(17, 22)):
            row = table5.rows[r_idx]
            row.cells[3].text = "X"
            apply_calibri_9(row.cells[3])
            for p in row.cells[3].paragraphs:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 2. Rellenar fecha de fin de la ejecución (Tabla 5, Fila 1, Celda 7)
        table5.rows[1].cells[7].text = execution_date_str
        apply_calibri_9(table5.rows[1].cells[7])
        for p in table5.rows[1].cells[7].paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 3. Rellenar Número de visitas en Tabla 5, Fila 1, Celda 10
        table5.rows[1].cells[10].text = "3"
        apply_calibri_9(table5.rows[1].cells[10])
        for p in table5.rows[1].cells[10].paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 4. Rellenar Tablas 6 (co-formador) y 7 (instructor) via inferencia JSON.
        #    SÍ rellenar Tabla 8 (aprendiz).
        table8 = doc.tables[8]
        table8.rows[2].cells[1].text = f"{actas_data['actas']['momento_3']['observaciones_aprendiz']}"
        apply_calibri_9(table8.rows[2].cells[1])

        # Tabla 6: Retroalimentación entre co-formador (inferencia)
        coformador_m3 = actas_data.get('actas', {}).get('momento_3', {}).get('observaciones_coformador', '')
        if coformador_m3:
            table6 = doc.tables[6]
            table6.rows[1].cells[1].text = coformador_m3
            apply_calibri_9(table6.rows[1].cells[1])
            print("        [Momento 3] Retroalimentación co-formador aplicada (Tabla 6).")

        # Tabla 7: Retroalimentación instructor de seguimiento (inferencia)
        instructor_m3 = actas_data.get('actas', {}).get('momento_3', {}).get('observaciones_instructor', '')
        if instructor_m3:
            table7 = doc.tables[7]
            table7.rows[1].cells[1].text = instructor_m3
            apply_calibri_9(table7.rows[1].cells[1])
            print("        [Momento 3] Retroalimentación instructor aplicada (Tabla 7).")

        # 5. Compromisos de mejora en Tabla 5, columna "Observaciones / Compromisos de mejora"
        compromisos = actas_data['actas']['momento_3'].get('compromisos_mejora', '')
        if compromisos:
            table5.rows[6].cells[6].text = compromisos
            apply_calibri_9(table5.rows[6].cells[6])
            table5.rows[17].cells[6].text = compromisos
            apply_calibri_9(table5.rows[17].cells[6])

        # 6. Marcar juicio de evaluación (búsqueda por texto)
        for idx in range(len(doc.paragraphs)):
            if "Juicio de evaluación" in doc.paragraphs[idx].text:
                doc.paragraphs[idx].text = "Juicio de evaluación de la etapa productiva (seleccione):     Aprobado [X]   No aprobado [ ]"
                apply_calibri_9(doc.paragraphs[idx])
                break

        # 7. Fecha de diligenciamiento en párrafo P33
        doc.paragraphs[33].text = f"\tEl momento 3 – Evaluación se llevó a cabo en la ciudad: Bogotá D.C. con fecha de diligenciamiento: {execution_date_str} de forma presencial ___ o virtual   X "
        apply_calibri_9(doc.paragraphs[33])

        print("        [Momento 3] Tabla 5 rellenada, tabla 6 (co-formador) y tabla 7 (instructor) actualizadas, tabla 8 (aprendiz) actualizada, "
              "compromisos de mejora agregados, juicio de evaluación marcado Aprobado, párrafo P33 actualizado.")

    doc.save(output_word_path)
    print(f"[Word] Guardado exitosamente: {output_word_path}")
    return output_word_path


# --- Estado de actas enviadas ---
_DEFAULT_ACTAS_STATE = {
    "momento_1": {"enviada": False, "fecha_envio": None},
    "momento_2": {"enviada": False, "fecha_envio": None},
    "momento_3": {"enviada": False, "fecha_envio": None},
}
ACTAS_STATE_PATH = os.path.join(WORK_DIR, "actas_enviadas.json")


def load_actas_state(state_path=None):
    if state_path is None:
        state_path = ACTAS_STATE_PATH
    if not os.path.exists(state_path):
        return {k: dict(v) for k, v in _DEFAULT_ACTAS_STATE.items()}
    try:
        with open(state_path, 'r', encoding='utf-8') as f:
            state = json.load(f)
    except (json.JSONDecodeError, OSError) as e:
        print(f"[Warning] Error al leer {state_path}: {e}", file=sys.stderr)
        return {k: dict(v) for k, v in _DEFAULT_ACTAS_STATE.items()}
    for key in ("momento_1", "momento_2", "momento_3"):
        if key not in state:
            state[key] = dict(_DEFAULT_ACTAS_STATE[key])
    return state


def save_actas_state(state, state_path=None):
    if state_path is None:
        state_path = ACTAS_STATE_PATH
    parent = os.path.dirname(state_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    with open(state_path, 'w', encoding='utf-8') as f:
        json.dump(state, f, indent=2, ensure_ascii=False)
        f.write('\n')


def mark_acta_enviada(state, moment, fecha_iso=None):
    moment_key = f"momento_{moment}"
    if fecha_iso is None:
        fecha_iso = datetime.date.today().isoformat()
    new_state = dict(state)
    new_state[moment_key] = {"enviada": True, "fecha_envio": fecha_iso}
    return new_state


def reset_actas_state(state):
    return {k: dict(v) for k, v in _DEFAULT_ACTAS_STATE.items()}


def main():
    parser = argparse.ArgumentParser(description="Automatización de documentos SENA (Excel y Word).")
    parser.add_argument("--date", type=str, default=None, help="Forzar fecha de ejecución (YYYY-MM-DD).")
    parser.add_argument(
        "--force-moment",
        type=int,
        choices=[1, 2, 3],
        default=None,
        help=(
            "Forzar el diligenciamiento de actas (Word). Acepta 1 (Planeacion), 2 (Seguimiento), "
            "3 (Evaluacion). Por defecto la fecha de ejecucion determina el momento via la ventana. "
            "ADVERTENCIA: --force-moment N activa el catch-up — si hay momentos anteriores (<= N) "
            "no enviados, tambien se enviaran en la misma ejecucion. Use --force-moment solo "
            "cuando quiera reenviar un momento especifico (revise contexto/actas_enviadas.json antes)."
        ),
    )
    parser.add_argument("--dry-run", action="store_true", help="Analizar markdown y memoria sin escribir archivos.")
    parser.add_argument("--no-email", action="store_true", help="No enviar correo electrónico al finalizar.")
    parser.add_argument(
        "--ai-mode",
        action="store_true",
        help="Modo agente: valida entorno y emite JSON estructurado si falta .env (no usa prompts).",
    )
    parser.add_argument(
        "--skip-env-check",
        action="store_true",
        help="Saltar la validacion del .env al inicio (util solo para tests internos).",
    )
    args = parser.parse_args()

    # 0. Validar variables de entorno al inicio.
    # Si falta algo, imprime error y sale. En modo agente imprime JSON estructurado.
    if not args.skip_env_check:
        if not env_validator.require_env(mode="ai" if args.ai_mode else "human"):
            if args.ai_mode:
                sys.exit(1)
            else:
                print(
                    "\n[Diligenciar] Configure el .env y vuelva a ejecutar.\n"
                    "              O use --skip-env-check solo si esta ejecutando tests internos.\n"
                    "              Wizard: python scripts/setup.py",
                    file=sys.stderr,
                )
                sys.exit(1)

    # Decidir envío de email: por defecto se envía, a menos que se indique --no-email o --dry-run
    send_email = not args.no_email and not args.dry_run
    if args.dry_run:
        print("[Dry Run] No se enviará correo electrónico (modo simulación).")
    elif args.no_email:
        print("[Email] Envío de correo desactivado por --no-email.")

    # 1. Resolver fecha de ejecución
    if args.date:
        exec_date = datetime.datetime.strptime(args.date, "%Y-%m-%d").date()
    else:
        exec_date = datetime.date.today()

    exec_date_str = exec_date.strftime("%d/%m/%Y")
    exec_date_folder = exec_date.strftime("%Y-%m-%d")

    print(f"=== INICIANDO AUTOMATIZACIÓN SENA ===")
    print(f"Fecha de ejecución: {exec_date} ({exec_date_str})")
    print(f"Modo email: {'HABILITADO' if send_email else 'DESHABILITADO'}")

    # 2. Leer archivos de entrada
    memory_data = parse_memory_descriptions()

    # 3. Encontrar TODAS las bitácoras pendientes (procesamiento por lote)
    pending_nums = get_all_undiligenced_bitacoras()
    all_generated_files = []
    last_bitacora_data = None
    bitacora_output_dir = None

    if not pending_nums:
        print("[Estado] No hay más bitácoras pendientes en historico_actividades.md.")
    else:
        print(f"[Estado] Bitácoras pendientes identificadas: {pending_nums}")

    for pending_num in pending_nums:
        print(f"\n--- Procesando Bitácora {pending_num} ---")

        bitacora_data = None
        for b in memory_data['bitacoras']:
            if b['numero'] == pending_num:
                bitacora_data = b
                last_bitacora_data = b
                break

        if bitacora_data is None:
            print(f"[Error] No se encontraron datos para la Bitácora {pending_num} en memory_descriptions.md!")
            continue

        if args.dry_run:
            print(f"[Dry Run] Se procesaría la Bitácora {pending_num} período {bitacora_data['fecha_inicio']} - {bitacora_data['fecha_fin']}.")
        else:
            # Determinar carpeta de salida para bitácora
            bitacora_output_dir = os.path.join(
                OUTPUT_DIR, f"bitacora{pending_num}-{exec_date_folder}"
            )
            os.makedirs(bitacora_output_dir, exist_ok=True)

            # Rellenar Excel
            excel_path = fill_excel_bitacora(pending_num, bitacora_data, exec_date, output_dir=bitacora_output_dir)
            all_generated_files.append(excel_path)

            # Marcar en markdown
            mark_bitacora_as_diligenced(pending_num)

    # 4. Procesar actas de Word (catch-up: enviar todas las pendientes <= momento actual)
    # 4.1. Resolver fechas configurables (.env) con defaults hardcodeados
    def _parse_env_iso_date(key, default):
        """Lee una fecha YYYY-MM-DD de .env. Si falta o es inválida, retorna default y avisa."""
        raw = os.getenv(key, "").strip()
        if not raw:
            return default
        try:
            return datetime.datetime.strptime(raw, "%Y-%m-%d").date()
        except ValueError:
            print(f"[Advertencia] {key}='{raw}' no es YYYY-MM-DD válido. Usando default {default}.")
            return default

    m1_date = _parse_env_iso_date("ACTA_M1_FECHA", datetime.date(2026, 4, 8))
    m2_date = _parse_env_iso_date("ACTA_M2_FECHA", datetime.date(2026, 6, 22))
    m3_date = _parse_env_iso_date("ACTA_M3_FECHA", datetime.date(2026, 10, 7))
    try:
        ventana = int(os.getenv("ACTA_VENTANA_DIAS", "7").strip())
    except ValueError:
        print("[Advertencia] ACTA_VENTANA_DIAS no es entero válido. Usando 7.")
        ventana = 7

    # 4.2. Determinar "momento actual" según la fecha de ejecución
    current_moment = None
    if abs((exec_date - m1_date).days) <= ventana:
        current_moment = 1
    elif abs((exec_date - m2_date).days) <= ventana:
        current_moment = 2
    elif abs((exec_date - m3_date).days) <= ventana:
        current_moment = 3

    # --force-moment sobrescribe current_moment (ignora ventana de fecha).
    # NOTA: el catch-up sigue activo — si N=3 y M1 y M2 no han sido enviadas,
    # se generaran las TRES actas (M1, M2, M3) en una sola ejecucion.
    if args.force_moment:
        current_moment = args.force_moment
        if current_moment not in (1, 2, 3):
            raise ValueError(f"--force-moment debe ser 1, 2 o 3. Recibido: {current_moment}")

    # 4.3. Cargar estado y calcular momentos a enviar
    state = load_actas_state()
    target_moments = []
    if current_moment is not None:
        for m in (1, 2, 3):
            if m <= current_moment and not state[f"momento_{m}"]["enviada"]:
                target_moments.append(m)
    else:
        # Sin current_moment: solo aplica si --force-moment fue usado
        if args.force_moment:
            target_moments = [args.force_moment]

    if not target_moments:
        print("[Estado] No hay actas pendientes para enviar.")

    # 4.4. Generar cada momento pendiente
    word_paths = []
    word_output_dirs = []
    for moment in target_moments:
        if args.dry_run:
            print(f"[Dry Run] Se diligenciaría el Acta de Word para el Momento {moment}.")
            continue
        # Carpeta de salida
        if pending_nums and bitacora_output_dir is not None:
            # Reusar la última carpeta de bitácora
            moment_output_dir = bitacora_output_dir
        else:
            moment_output_dir = os.path.join(
                OUTPUT_DIR, f"acta-momento{moment}-{exec_date_folder}"
            )
            os.makedirs(moment_output_dir, exist_ok=True)
        word_path = process_word_actas(moment, memory_data, exec_date_str, output_dir=moment_output_dir)
        all_generated_files.append(word_path)
        word_paths.append(word_path)
        word_output_dirs.append(moment_output_dir)
        # Marcar como enviada en el estado
        state = mark_acta_enviada(state, moment, fecha_iso=exec_date.isoformat())

    # Persistir estado de actas enviadas
    save_actas_state(state)
    print(f"[Estado] Estado de actas enviadas persistido en {ACTAS_STATE_PATH}")

    # 5. Escribir archivo de log en la carpeta de salida correspondiente
    if pending_nums and last_bitacora_data is not None:
        log_dir = bitacora_output_dir
    elif target_moments and word_output_dirs and not args.dry_run:
        log_dir = word_output_dirs[-1]
    else:
        log_dir = None

    if log_dir and not args.dry_run:
        log_path = os.path.join(log_dir, "ejecucion.log")
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write(f"Fecha de ejecución: {exec_date} ({exec_date_str})\n")
            f.write(f"Bitácoras procesadas: {pending_nums if pending_nums else 'Ninguna'}\n")
            if pending_nums and last_bitacora_data:
                primera = memory_data['bitacoras'][0]
                ultima = last_bitacora_data
                f.write(f"Período total: {primera['fecha_inicio']} - {ultima['fecha_fin']}\n")
            f.write(f"Acta Word Momentos: {target_moments if target_moments else 'N/A'}\n")
            f.write(f"Estado: COMPLETADO\n")
        print(f"[Log] Resumen guardado en {log_path}")

    # 6. Envío de email si está habilitado
    if send_email and not args.dry_run:
        from email_module import (
            enviar_email,
            construir_asunto,
            construir_cuerpo,
            reintentar_envio_manual,
        )

        modo = os.getenv("EMAIL_MODO", "pruebas")
        if modo == "produccion":
            destinatario = os.getenv("EMAIL_DESTINATARIO_PRODUCCION")
        else:
            destinatario = os.getenv("EMAIL_DESTINATARIO_PRUEBAS")

        if not destinatario:
            raise ValueError(
                f"EMAIL_DESTINATARIO_{'PRODUCCION' if modo == 'produccion' else 'PRUEBAS'}"
                f" no esta configurado en .env. Agregalo antes de continuar."
            )

        cc = os.getenv("EMAIL_CC")
        if not cc:
            raise ValueError("EMAIL_CC no esta configurado en .env. Agregalo antes de continuar.")

        # Recopilar todos los archivos generados (xlsx, docx) desde los directorios de salida
        adjuntos = all_generated_files[:]

        # Construir lista de bitácoras procesadas con sus períodos
        bitacoras_info = []
        if pending_nums:
            for b in memory_data['bitacoras']:
                if b['numero'] in pending_nums:
                    bitacoras_info.append({
                        'numero': b['numero'],
                        'fecha_inicio': b['fecha_inicio'],
                        'fecha_fin': b['fecha_fin'],
                    })

        asunto = construir_asunto(
            bitacoras_info, target_moments
        )
        cuerpo = construir_cuerpo(bitacoras_info, target_moments)

        exito, mensaje = enviar_email(
            destinatario=destinatario,
            cc=cc,
            asunto=asunto,
            cuerpo=cuerpo,
            adjuntos=adjuntos,
            intentos=3,
        )

        if exito:
            print(f"[Email] {mensaje}")
        else:
            print(f"[Email] {mensaje}")
            if log_dir:
                reintentar_envio_manual(bitacoras_info, log_dir, adjuntos)

    elif send_email and args.dry_run:
        print(
            "[Dry Run] Se habría enviado un correo electrónico"
            " con los documentos adjuntos."
        )

    # 7. Checklist de campos que requieren diligenciamiento/firma manual
    if target_moments and not args.dry_run:
        print()
        print("=" * 70)
        print("[CHECKLIST MANUAL] Campos que requieren diligenciamiento/firma humana:")
        print("=" * 70)
        # Definir checklist por momento
        CHECKLIST = {
            1: [
                "Tabla 2 R3: Enlace de grabación del momento 1 (si aplica)",
                "Tabla 2 R11: Firmas (aprendiz, instructor seguimiento, ente co-formador)",
            ],
            2: [
                "Párrafo 14: Observaciones del instructor de seguimiento (texto en blanco, lo llena el instructor)",
                "Tabla 3 R2: Enlace de grabación del momento 2 (si aplica)",
                "Tabla 4: Firmas (aprendiz, instructor seguimiento, ente co-formador)",
            ],
            3: [
                "Tabla 5 R2: Enlace de grabación del momento 3 (si aplica)",
                "Tabla 9: Firmas (aprendiz, instructor seguimiento, ente co-formador)",
            ],
        }
        # Checklist general (Tabla 1) - aplica a TODOS los momentos
        print("GENERAL (Tabla 1 - aplica a todos los momentos):")
        print("  - Regional, Centro de formación, Programa de formación")
        print("  - Datos del aprendiz: documento, teléfono, dirección, correos")
        print("  - Datos del instructor: teléfono, correo institucional")
        print("  - Datos del ente co-formador: dirección, NIT, correo, jefe, cargo, teléfono")
        for m in target_moments:
            print(f"\nMOMENTO {m}:")
            for item in CHECKLIST.get(m, []):
                print(f"  - {item}")
        print("=" * 70)
        # Guardar checklist en cada carpeta de salida de acta
        for d in word_output_dirs:
            checklist_path = os.path.join(d, "checklist_manual.txt")
            try:
                with open(checklist_path, 'w', encoding='utf-8') as f:
                    f.write("CHECKLIST MANUAL - Campos que requieren diligenciamiento/firma humana\n")
                    f.write("=" * 70 + "\n\n")
                    f.write("GENERAL (Tabla 1 - aplica a todos los momentos):\n")
                    f.write("  - Regional, Centro de formación, Programa de formación\n")
                    f.write("  - Datos del aprendiz: documento, teléfono, dirección, correos\n")
                    f.write("  - Datos del instructor: teléfono, correo institucional\n")
                    f.write("  - Datos del ente co-formador: dirección, NIT, correo, jefe, cargo, teléfono\n\n")
                    for m in target_moments:
                        f.write(f"MOMENTO {m}:\n")
                        for item in CHECKLIST.get(m, []):
                            f.write(f"  - {item}\n")
                        f.write("\n")
                print(f"[Checklist] Guardado en {checklist_path}")
            except OSError as e:
                print(f"[Checklist] No se pudo escribir {checklist_path}: {e}")

    print(f"=== AUTOMATIZACIÓN SENA FINALIZADA CON ÉXITO ===")


if __name__ == "__main__":
    main()
