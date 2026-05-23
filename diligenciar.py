import os
import shutil
import json
import datetime
import argparse
import re
import openpyxl
from openpyxl.styles import Alignment, Border, Side, PatternFill, Font
import docx
from docx.shared import Pt
from dotenv import load_dotenv

load_dotenv()

# Caminos absolutos a los archivos de trabajo
WORK_DIR = "/home/eivorkinkest/Documentos/Docs_SENA/DocsOriginales"
EXCEL_PATH = os.path.join(WORK_DIR, "BitacoraMQuiazua1.xlsx")
TEMPLATE_EXCEL_PATH = os.path.join(WORK_DIR, "BitacoraMQuiazua_template.xlsx")
WORD_PATH = os.path.join(WORK_DIR, "Actas-Inicio-Medio-Final.docx")
TEMPLATE_WORD_PATH = os.path.join(WORK_DIR, "Actas-Inicio-Medio-Final_template.docx")
HISTORICO_PATH = os.path.join(WORK_DIR, "historico_actividades.md")
MEMORY_PATH = os.path.join(WORK_DIR, "memory_descriptions.md")

# Constantes de layout
EXCEL_ACT_START_ROW = 40
EXCEL_ACT_ROW_SPAN = 2
EXCEL_MAX_ACTIVITIES = 7
EXCEL_ACT_END_ROW = EXCEL_ACT_START_ROW + (EXCEL_MAX_ACTIVITIES * EXCEL_ACT_ROW_SPAN) - 1  # 53


def parse_memory_descriptions():
    """
    Lee y parsea el bloque JSON contenido dentro del archivo memory_descriptions.md
    """
    if not os.path.exists(MEMORY_PATH):
        raise FileNotFoundError(f"No se encontró el archivo de memoria en {MEMORY_PATH}")

    with open(MEMORY_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extraer el bloque JSON usando regex
    match = re.search(r"```json\s*(.*?)\s*```", content, re.DOTALL)
    if not match:
        raise ValueError("No se encontró el bloque ```json ... ``` en memory_descriptions.md")

    json_str = match.group(1)
    return json.loads(json_str)


def get_next_undiligenced_bitacora():
    """
    Lee historico_actividades.md e identifica el número de la primera bitácora
    que no esté marcada como [DILIGENCIADA].
    """
    if not os.path.exists(HISTORICO_PATH):
        raise FileNotFoundError(f"No se encontró el histórico en {HISTORICO_PATH}")

    with open(HISTORICO_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        match = re.search(r"##\s*[Bb]itacora\s*numero\s*(\d+)", line)
        if match:
            bitacora_num = int(match.group(1))
            if "[DILIGENCIADA]" not in line:
                return bitacora_num
    return None


def mark_bitacora_as_diligenced(num):
    """
    Modifica historico_actividades.md para añadir la etiqueta [DILIGENCIADA]
    al encabezado de la bitácora número 'num'.
    """
    if not os.path.exists(HISTORICO_PATH):
        return

    with open(HISTORICO_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    pattern = rf"(##\s*[Bb]itacora\s*numero\s*{num}\s*-[^\n]+)"

    def replacer(match):
        line = match.group(1)
        if "[DILIGENCIADA]" not in line:
            return f"{line.strip()} [DILIGENCIADA]"
        return line

    new_content, count = re.subn(pattern, replacer, content)
    if count > 0:
        with open(HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(new_content)
        print(f"[Estado] historico_actividades.md actualizado: Bitácora {num} marcada como [DILIGENCIADA].")


def copy_style(src_cell, dest_cell):
    """
    Copia el estilo visual de una celda origen a una celda destino
    """
    if src_cell.has_style:
        dest_cell.font = Font(
            name=src_cell.font.name,
            size=src_cell.font.size,
            bold=src_cell.font.bold,
            italic=src_cell.font.italic,
            color=src_cell.font.color,
            underline=src_cell.font.underline
        )
        dest_cell.border = Border(
            left=src_cell.border.left,
            right=src_cell.border.right,
            top=src_cell.border.top,
            bottom=src_cell.border.bottom
        )
        dest_cell.fill = PatternFill(
            fill_type=src_cell.fill.fill_type,
            start_color=src_cell.fill.start_color,
            end_color=src_cell.fill.end_color
        )
        dest_cell.alignment = Alignment(
            horizontal=src_cell.alignment.horizontal,
            vertical=src_cell.alignment.vertical,
            wrap_text=src_cell.alignment.wrap_text
        )
        dest_cell.number_format = src_cell.number_format


def fill_excel_bitacora(bitacora_num, bitacora_data, execution_date, output_dir=None):
    """
    Rellena la bitácora en Excel usando openpyxl.
    Respalda la plantilla original si no se ha hecho aún, y guarda el resultado.
    Cada actividad ocupa 2 filas con merge vertical en B, C, D, E, F.
    Máximo 7 actividades (filas 40-53).
    """
    # 1. Asegurar el respaldo de la plantilla limpia
    if not os.path.exists(TEMPLATE_EXCEL_PATH):
        if not os.path.exists(EXCEL_PATH):
            raise FileNotFoundError(f"No se encontró la plantilla original en {EXCEL_PATH}")
        shutil.copy2(EXCEL_PATH, TEMPLATE_EXCEL_PATH)
        print(f"[Template] Respaldo de plantilla creado en {TEMPLATE_EXCEL_PATH}")

    # 2. Definir la carpeta de salida
    if output_dir is None:
        date_str = execution_date.strftime("%Y-%m-%d")
        output_dir = os.path.join(WORK_DIR, "output", f"bitacora{bitacora_num}-{date_str}")
    os.makedirs(output_dir, exist_ok=True)

    # 3. Definir el archivo de salida
    output_filename = f"BitacoraMQuiazua{bitacora_num}.xlsx"
    output_path = os.path.join(output_dir, output_filename)

    # 4. Copiar de la plantilla limpia al archivo final
    shutil.copy2(TEMPLATE_EXCEL_PATH, output_path)

    # 5. Abrir y diligenciar
    wb = openpyxl.load_workbook(output_path)
    sheet = wb['GFPI-F-147-Formato Bitácora']

    # Metadatos del encabezado
    sheet['E17'] = bitacora_num
    sheet['F17'] = f"DESDE {bitacora_data['fecha_inicio']}"
    sheet['G17'] = f"HASTA {bitacora_data['fecha_fin']}"
    sheet['F67'] = execution_date

    print(f"[Excel] Diligenciando {output_filename} para Bitácora {bitacora_num}:")
    print(f"        Período: {bitacora_data['fecha_inicio']} - {bitacora_data['fecha_fin']}")
    print(f"        Carpeta: {output_dir}")

    # Limpiar y rellenar actividades (filas 40 a 53)
    # Descombinar cualquier combinación previa en las filas de actividad
    for r in range(EXCEL_ACT_START_ROW, EXCEL_ACT_END_ROW + 1):
        for col in ['B', 'C', 'D', 'E', 'F', 'G']:
            sheet[f"{col}{r}"].value = None
        for rng in list(sheet.merged_cells.ranges):
            if rng.min_row == r and rng.min_col in (2, 4, 5, 6):
                sheet.unmerge_cells(rng.coord)

    # Copiar estilos base de la celda de referencia B40
    ref_cell_b = sheet['B40']
    ref_cell_d = sheet['D40']
    ref_cell_e = sheet['E40']
    ref_cell_f = sheet['F40']
    ref_cell_g = sheet['G40']

    activities = bitacora_data['actividades']
    max_activities_shown = min(len(activities), EXCEL_MAX_ACTIVITIES)

    if len(activities) > EXCEL_MAX_ACTIVITIES:
        print(f"[Warning] Hay {len(activities)} actividades pero el límite en Excel es {EXCEL_MAX_ACTIVITIES}. "
              f"Se mostrarán solo las primeras {EXCEL_MAX_ACTIVITIES}.")

    for idx in range(max_activities_shown):
        act = activities[idx]
        r = EXCEL_ACT_START_ROW + (idx * EXCEL_ACT_ROW_SPAN)
        r2 = r + 1  # segunda fila del par

        # Escribir la descripción en B{r} y combinar B{r}:B{r2}
        sheet[f"B{r}"] = act['descripcion']
        sheet.merge_cells(start_row=r, start_column=2, end_row=r2, end_column=3)
        copy_style(ref_cell_b, sheet[f"B{r}"])
        sheet[f"B{r}"].alignment = Alignment(wrap_text=True, vertical="center", horizontal="left")
        copy_style(ref_cell_b, sheet[f"C{r}"])

        # Merge vertical para D, E, F (de r a r2)
        sheet[f"D{r}"] = act['fecha_inicio']
        sheet.merge_cells(start_row=r, start_column=4, end_row=r2, end_column=4)
        copy_style(ref_cell_d, sheet[f"D{r}"])

        sheet[f"E{r}"] = act['fecha_fin']
        sheet.merge_cells(start_row=r, start_column=5, end_row=r2, end_column=5)
        copy_style(ref_cell_e, sheet[f"E{r}"])

        sheet[f"F{r}"] = act['evidencia']
        sheet.merge_cells(start_row=r, start_column=6, end_row=r2, end_column=6)
        copy_style(ref_cell_f, sheet[f"F{r}"])
        sheet[f"F{r}"].alignment = Alignment(wrap_text=True, vertical="center", horizontal="left")

        # Celda G vacía con estilo
        sheet[f"G{r}"] = ""
        copy_style(ref_cell_g, sheet[f"G{r}"])

        print(f"        Filas {r}-{r2}: {act['descripcion'][:40]}... | {act['fecha_inicio']} al {act['fecha_fin']}")

    wb.save(output_path)
    print(f"[Excel] Grabado exitosamente: {output_path}")
    return output_path


def apply_calibri_9(element):
    """
    Aplica tipografía Calibri tamaño 9 a un párrafo o celda de tabla de Word.
    Afecta todos los runs existentes del elemento.
    """
    if isinstance(element, docx.text.paragraph.Paragraph):
        for run in element.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
    elif isinstance(element, docx.table._Cell):
        for paragraph in element.paragraphs:
            apply_calibri_9(paragraph)


def process_word_actas(moment, actas_data, execution_date_str, output_dir=None):
    """
    Rellena el momento correspondiente (2 o 3) en el Word.
    """
    if moment not in [2, 3]:
        print(f"[Word] Momento {moment} inválido para actas. Ignorando.")
        return

    # --- Respaldo de plantilla Word ---
    if not os.path.exists(TEMPLATE_WORD_PATH):
        if not os.path.exists(WORD_PATH):
            raise FileNotFoundError(f"No se encontró la plantilla original de Word en {WORD_PATH}")
        shutil.copy2(WORD_PATH, TEMPLATE_WORD_PATH)
        print(f"[Template Word] Respaldo de plantilla creado en {TEMPLATE_WORD_PATH}")

    # --- Determinar ruta de salida ---
    if output_dir is None:
        output_dir = os.path.join(WORK_DIR, "output")
    os.makedirs(output_dir, exist_ok=True)
    output_word_path = os.path.join(output_dir, "Actas-Inicio-Medio-Final.docx")

    # --- Copiar desde el template limpio ---
    shutil.copy2(TEMPLATE_WORD_PATH, output_word_path)

    print(f"[Word] Diligenciando Acta para Momento {moment} (Fecha ejecución: {execution_date_str}):")
    print(f"        Archivo: {output_word_path}")

    doc = docx.Document(output_word_path)

    if moment == 2:
        # --- MOMENTO 2: SEGUIMIENTO ---
        # 1. Rellenar marcas X en Tabla 3 (Seguimiento).
        table3 = doc.tables[3]
        for r_idx in list(range(6, 14)) + list(range(17, 22)):
            row = table3.rows[r_idx]
            row.cells[2].text = "X"
            apply_calibri_9(row.cells[2])
            for p in row.cells[2].paragraphs:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 2. Rellenar fecha de momento en Tabla 3, Fila 1, Celda 8
        table3.rows[1].cells[8].text = execution_date_str
        apply_calibri_9(table3.rows[1].cells[8])
        for p in table3.rows[1].cells[8].paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 3. Rellenar SOLO observaciones del aprendiz (P17).
        #    NO rellenar P14 (instructor seguimiento) ni P20 (co-formador)
        doc.paragraphs[17].text = f"  {actas_data['actas']['momento_2']['observaciones_aprendiz']}"
        apply_calibri_9(doc.paragraphs[17])

        # 4. Compromisos de mejora en Tabla 3, columna "Observaciones / Compromisos de mejora"
        compromisos = actas_data['actas']['momento_2'].get('compromisos_mejora', '')
        if compromisos:
            # Escribir en la primera fila técnica (row 6, col 6) y actitudinal (row 17, col 6)
            table3.rows[6].cells[6].text = compromisos
            apply_calibri_9(table3.rows[6].cells[6])
            table3.rows[17].cells[6].text = compromisos
            apply_calibri_9(table3.rows[17].cells[6])

        # 5. Fecha de diligenciamiento en párrafo P22
        doc.paragraphs[22].text = f"Ciudad: Bogotá D.C.  y fecha de diligenciamiento: {execution_date_str} de forma presencial ___ o virtual   X "
        apply_calibri_9(doc.paragraphs[22])

        print("        [Momento 2] Tabla 3 rellenada, observaciones del aprendiz insertadas, "
              "compromisos de mejora agregados, párrafo P22 actualizado.")

    elif moment == 3:
        # --- MOMENTO 3: EVALUACIÓN FINAL ---
        # 1. Rellenar marcas X en Tabla 5 (Evaluación).
        table5 = doc.tables[5]
        for r_idx in list(range(6, 14)) + list(range(17, 22)):
            row = table5.rows[r_idx]
            row.cells[3].text = "X"
            apply_calibri_9(row.cells[3])
            for p in row.cells[3].paragraphs:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 2. Rellenar fecha de fin de la ejecución (Tabla 5, Fila 1, Celda 7)
        table5.rows[1].cells[7].text = execution_date_str
        apply_calibri_9(table5.rows[1].cells[7])
        for p in table5.rows[1].cells[7].paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 3. Rellenar Número de visitas en Tabla 5, Fila 1, Celda 10
        table5.rows[1].cells[10].text = "3"
        apply_calibri_9(table5.rows[1].cells[10])
        for p in table5.rows[1].cells[10].paragraphs:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # 4. NO rellenar Tablas 6 (co-formador) ni 7 (instructor).
        #    SÍ rellenar Tabla 8 (aprendiz).
        table8 = doc.tables[8]
        table8.rows[1].cells[1].text = f"He fortalecido mis bases de ingeniería y he aprendido a desarrollar con altos estándares de calidad, logrando cumplir mis objetivos."
        apply_calibri_9(table8.rows[1].cells[1])
        table8.rows[2].cells[1].text = f"{actas_data['actas']['momento_3']['observaciones_aprendiz']}"
        apply_calibri_9(table8.rows[2].cells[1])

        # 5. Compromisos de mejora en Tabla 5, columna "Observaciones / Compromisos de mejora"
        compromisos = actas_data['actas']['momento_3'].get('compromisos_mejora', '')
        if compromisos:
            table5.rows[6].cells[6].text = compromisos
            apply_calibri_9(table5.rows[6].cells[6])
            table5.rows[17].cells[6].text = compromisos
            apply_calibri_9(table5.rows[17].cells[6])

        # 6. Marcar juicio de evaluación (búsqueda por texto)
        for idx in range(len(doc.paragraphs)):
            if "Juicio de evaluación" in doc.paragraphs[idx].text:
                doc.paragraphs[idx].text = "Juicio de evaluación de la etapa productiva (seleccione):     Aprobado [X]   No aprobado [ ]"
                apply_calibri_9(doc.paragraphs[idx])
                break

        # 7. Fecha de diligenciamiento en párrafo P33
        doc.paragraphs[33].text = f"\tEl momento 3 – Evaluación se llevó a cabo en la ciudad: Bogotá D.C. con fecha de diligenciamiento: {execution_date_str} de forma presencial ___ o virtual   X "
        apply_calibri_9(doc.paragraphs[33])

        print("        [Momento 3] Tabla 5 rellenada, tabla 8 actualizada (aprendiz), "
              "compromisos de mejora agregados, juicio de evaluación marcado Aprobado, párrafo P33 actualizado.")

    doc.save(output_word_path)
    print(f"[Word] Guardado exitosamente: {output_word_path}")
    return output_word_path


def main():
    parser = argparse.ArgumentParser(description="Automatización de documentos SENA (Excel y Word).")
    parser.add_argument("--date", type=str, default=None, help="Forzar fecha de ejecución (YYYY-MM-DD).")
    parser.add_argument("--force-moment", type=int, choices=[2, 3], default=None, help="Forzar el diligenciamiento de un acta (Word).")
    parser.add_argument("--dry-run", action="store_true", help="Analizar markdown y memoria sin escribir archivos.")
    args = parser.parse_args()

    # Preguntar autorización de email
    send_email = False
    if not args.dry_run:
        from email_module import preguntar_envio_email
        send_email = preguntar_envio_email()
    else:
        print("[Dry Run] Se preguntaría si desea enviar correo electrónico al finalizar.")

    # 1. Resolver fecha de ejecución
    if args.date:
        exec_date = datetime.datetime.strptime(args.date, "%Y-%m-%d").date()
    else:
        exec_date = datetime.date.today()

    exec_date_str = exec_date.strftime("%d/%m/%Y")
    exec_date_folder = exec_date.strftime("%Y-%m-%d")

    print(f"=== INICIANDO AUTOMATIZACIÓN SENA ===")
    print(f"Fecha de ejecución: {exec_date} ({exec_date_str})")

    # 2. Leer archivos de entrada
    memory_data = parse_memory_descriptions()

    # 3. Encontrar la siguiente bitácora pendiente en el markdown
    pending_num = get_next_undiligenced_bitacora()
    if pending_num is None:
        print("[Estado] No hay más bitácoras pendientes en historico_actividades.md.")
    else:
        print(f"[Estado] Bitácora pendiente identificada: Número {pending_num}")

        bitacora_data = None
        for b in memory_data['bitacoras']:
            if b['numero'] == pending_num:
                bitacora_data = b
                break

        if bitacora_data is None:
            print(f"[Error] No se encontraron datos para la Bitácora {pending_num} en memory_descriptions.md!")
        else:
            if args.dry_run:
                print(f"[Dry Run] Se procesaría la Bitácora {pending_num} período {bitacora_data['fecha_inicio']} - {bitacora_data['fecha_fin']}.")
            else:
                # Determinar carpeta de salida para bitácora
                bitacora_output_dir = os.path.join(
                    WORK_DIR, "output", f"bitacora{pending_num}-{exec_date_folder}"
                )
                os.makedirs(bitacora_output_dir, exist_ok=True)

                # Rellenar Excel
                excel_path = fill_excel_bitacora(pending_num, bitacora_data, exec_date, output_dir=bitacora_output_dir)

                # Marcar en markdown
                mark_bitacora_as_diligenced(pending_num)

    # 4. Procesar actas de Word
    target_moment = None
    if args.force_moment:
        target_moment = args.force_moment
    else:
        m2_date = datetime.date(2026, 7, 8)
        m3_date = datetime.date(2026, 10, 8)

        diff_m2 = abs((exec_date - m2_date).days)
        diff_m3 = abs((exec_date - m3_date).days)

        if diff_m2 <= 7:
            target_moment = 2
        elif diff_m3 <= 7:
            target_moment = 3

    if target_moment:
        if args.dry_run:
            print(f"[Dry Run] Se diligenciaría el Acta de Word para el Momento {target_moment}.")
        else:
            # Determinar carpeta de salida para Word
            if pending_num is not None and bitacora_data is not None:
                # Si también hay bitácora, reusar la misma carpeta
                word_output_dir = bitacora_output_dir
            else:
                # Carpeta independiente para Word
                word_output_dir = os.path.join(
                    WORK_DIR, "output", f"acta-momento{target_moment}-{exec_date_folder}"
                )
                os.makedirs(word_output_dir, exist_ok=True)

            word_path = process_word_actas(target_moment, memory_data, exec_date_str, output_dir=word_output_dir)

    # 5. Escribir archivo de log en la carpeta de salida correspondiente
    if pending_num is not None and bitacora_data is not None:
        log_dir = bitacora_output_dir
    elif target_moment and not args.dry_run:
        log_dir = word_output_dir
    else:
        log_dir = None

    if log_dir and not args.dry_run:
        log_path = os.path.join(log_dir, "ejecucion.log")
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write(f"Fecha de ejecución: {exec_date} ({exec_date_str})\n")
            f.write(f"Bitácora procesada: {pending_num if pending_num else 'N/A'}\n")
            if pending_num and bitacora_data:
                f.write(f"Período: {bitacora_data['fecha_inicio']} - {bitacora_data['fecha_fin']}\n")
                f.write(f"Actividades: {len(bitacora_data['actividades'])}\n")
            f.write(f"Acta Word Momento: {target_moment if target_moment else 'N/A'}\n")
            f.write(f"Estado: COMPLETADO\n")
        print(f"[Log] Resumen guardado en {log_path}")

    # 6. Envío de email si autorizado
    if send_email and not args.dry_run:
        from email_module import (
            enviar_email,
            construir_asunto,
            construir_cuerpo,
            reintentar_envio_manual,
        )

        modo = os.getenv("EMAIL_MODO", "pruebas")
        if modo == "produccion":
            destinatario = os.getenv(
                "EMAIL_DESTINATARIO_PRODUCCION", "oiospina@sena.edu.co"
            )
        else:
            destinatario = os.getenv(
                "EMAIL_DESTINATARIO_PRUEBAS", "jmqo2015@gmail.com"
            )
        cc = os.getenv("EMAIL_CC", "eivorkinkest@gmail.com")

        fecha_inicio = (
            bitacora_data.get("fecha_inicio", "") if bitacora_data else ""
        )
        fecha_fin = (
            bitacora_data.get("fecha_fin", "") if bitacora_data else ""
        )

        asunto = construir_asunto(
            pending_num, fecha_inicio, fecha_fin, target_moment
        )
        cuerpo = construir_cuerpo(fecha_inicio, fecha_fin, target_moment)

        adjuntos = []
        if log_dir and os.path.exists(log_dir):
            for fname in os.listdir(log_dir):
                if fname.endswith((".xlsx", ".docx")):
                    adjuntos.append(os.path.join(log_dir, fname))

        exito, mensaje = enviar_email(
            destinatario=destinatario,
            cc=cc,
            asunto=asunto,
            cuerpo=cuerpo,
            adjuntos=adjuntos,
            intentos=3,
        )

        if exito:
            print(f"[Email] {mensaje}")
        else:
            print(f"[Email] {mensaje}")
            if log_dir:
                reintentar_envio_manual(pending_num, log_dir)

    elif send_email and args.dry_run:
        print(
            "[Dry Run] Se habría enviado un correo electrónico"
            " con los documentos adjuntos."
        )

    print(f"=== AUTOMATIZACIÓN SENA FINALIZADA CON ÉXITO ===")


if __name__ == "__main__":
    main()
