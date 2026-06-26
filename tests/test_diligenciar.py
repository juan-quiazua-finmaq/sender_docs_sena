import os
import sys
import shutil
import tempfile
import unittest
import json
import datetime
import warnings
import openpyxl
import docx
from docx.shared import Pt

# Asegurar que scripts/ está en sys.path para importar diligenciar
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))
import diligenciar
import email_module
import env_validator
from unittest.mock import patch


class TestDiligenciarAutomation(unittest.TestCase):
    
    def setUp(self):
        # Crear un directorio temporal para pruebas aisladas
        self.test_dir = tempfile.mkdtemp()
        
        # Guardar caminos originales y redirigir los de diligenciar.py al directorio temporal
        self.original_work_dir = diligenciar.WORK_DIR
        self.original_output_dir = diligenciar.OUTPUT_DIR
        self.original_excel_path = diligenciar.EXCEL_PATH
        self.original_template_excel_path = diligenciar.TEMPLATE_EXCEL_PATH
        self.original_word_path = diligenciar.WORD_PATH
        self.original_template_word_path = diligenciar.TEMPLATE_WORD_PATH
        self.original_historico_path = diligenciar.HISTORICO_PATH
        self.original_memory_path = diligenciar.MEMORY_PATH
        
        # Las plantillas ahora están en contexto/plantillas/ con nombres simplificados
        plantillas_dir = os.path.join(self.test_dir, "plantillas")
        os.makedirs(plantillas_dir, exist_ok=True)
        
        diligenciar.WORK_DIR = self.test_dir
        diligenciar.OUTPUT_DIR = os.path.join(self.test_dir, "output")
        diligenciar.EXCEL_PATH = os.path.join(plantillas_dir, "bitacoras.xlsx")
        diligenciar.TEMPLATE_EXCEL_PATH = os.path.join(plantillas_dir, "bitacoras.xlsx")
        diligenciar.WORD_PATH = os.path.join(plantillas_dir, "actas.docx")
        diligenciar.TEMPLATE_WORD_PATH = os.path.join(plantillas_dir, "actas.docx")
        diligenciar.HISTORICO_PATH = os.path.join(self.test_dir, "historico_actividades.md")
        diligenciar.MEMORY_PATH = os.path.join(self.test_dir, "memory_descriptions.md")
        
        self.original_actas_state_path = diligenciar.ACTAS_STATE_PATH
        diligenciar.ACTAS_STATE_PATH = os.path.join(self.test_dir, "actas_enviadas.json")
        
        # Copiar los archivos plantilla reales al directorio temporal para pruebas de integración reales
        if os.path.exists(self.original_excel_path):
            shutil.copy2(self.original_excel_path, diligenciar.EXCEL_PATH)
        if os.path.exists(self.original_word_path):
            shutil.copy2(self.original_word_path, diligenciar.WORD_PATH)
        if os.path.exists(self.original_historico_path):
            shutil.copy2(self.original_historico_path, diligenciar.HISTORICO_PATH)
        if os.path.exists(self.original_memory_path):
            shutil.copy2(self.original_memory_path, diligenciar.MEMORY_PATH)

    def tearDown(self):
        # Restaurar caminos originales
        diligenciar.WORK_DIR = self.original_work_dir
        diligenciar.OUTPUT_DIR = self.original_output_dir
        diligenciar.EXCEL_PATH = self.original_excel_path
        diligenciar.TEMPLATE_EXCEL_PATH = self.original_template_excel_path
        diligenciar.WORD_PATH = self.original_word_path
        diligenciar.TEMPLATE_WORD_PATH = self.original_template_word_path
        diligenciar.HISTORICO_PATH = self.original_historico_path
        diligenciar.MEMORY_PATH = self.original_memory_path
        diligenciar.ACTAS_STATE_PATH = self.original_actas_state_path
        
        # Eliminar el directorio temporal
        shutil.rmtree(self.test_dir)

    # =========================================================================
    # TESTS UNITARIOS (existentes)
    # =========================================================================

    def test_parse_memory_descriptions(self):
        """
        Verifica que la función parse_memory_descriptions extraiga el JSON correctamente
        """
        data = diligenciar.parse_memory_descriptions()
        self.assertIn("bitacoras", data)
        self.assertIn("actas", data)
        self.assertGreaterEqual(len(data["bitacoras"]), 1)
        self.assertEqual(data["bitacoras"][0]["numero"], 1)

    def test_get_next_undiligenced_bitacora(self):
        """
        Verifica que get_next_undiligenced_bitacora identifique correctamente la bitácora pendiente
        """
        mock_markdown = (
            "# Historico\n\n"
            "## Bitacora numero 1 - (08/04/2026 al 22/04/2026) [DILIGENCIADA]\n"
            "- Actividad 1\n"
            "## Bitacora numero 2 - (22/04/2026 al 06/05/2026)\n"
            "- Actividad 2\n"
            "## Bitacora numero 3 - (06/05/2026 al 20/05/2026)\n"
            "- Actividad 3\n"
        )
        with open(diligenciar.HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(mock_markdown)
            
        pending = diligenciar.get_next_undiligenced_bitacora()
        self.assertEqual(pending, 2)

    def test_mark_bitacora_as_diligenced(self):
        """
        Verifica que mark_bitacora_as_diligenced marque la bitácora correspondiente en el markdown
        """
        mock_markdown = (
            "# Historico\n\n"
            "## Bitacora numero 1 - (08/04/2026 al 22/04/2026)\n"
            "- Actividad 1\n"
        )
        with open(diligenciar.HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(mock_markdown)
            
        diligenciar.mark_bitacora_as_diligenced(1)
        
        with open(diligenciar.HISTORICO_PATH, 'r', encoding='utf-8') as f:
            content = f.read()
            
        self.assertIn("## Bitacora numero 1 - (08/04/2026 al 22/04/2026) [DILIGENCIADA]", content)

    # =========================================================================
    # TESTS UNITARIOS: get_all_undiligenced_bitacoras
    # =========================================================================

    def test_get_all_undiligenced_bitacoras_todas_diligenciadas(self):
        """Retorna lista vacía si todas las bitácoras están diligenciadas."""
        mock_markdown = (
            "# Historico\n\n"
            "## Bitacora numero 1 - (08/04/2026 al 22/04/2026) [DILIGENCIADA]\n"
            "- Actividad 1\n"
            "## Bitacora numero 2 - (22/04/2026 al 06/05/2026) [DILIGENCIADA]\n"
            "- Actividad 2\n"
        )
        with open(diligenciar.HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(mock_markdown)

        result = diligenciar.get_all_undiligenced_bitacoras()
        self.assertEqual(result, [])

    def test_get_all_undiligenced_bitacoras_multiples_pendientes(self):
        """Retorna los números de todas las bitácoras pendientes."""
        mock_markdown = (
            "# Historico\n\n"
            "## Bitacora numero 1 - (08/04/2026 al 22/04/2026)\n"
            "- Actividad 1\n"
            "## Bitacora numero 2 - (22/04/2026 al 06/05/2026) [DILIGENCIADA]\n"
            "- Actividad 2\n"
            "## Bitacora numero 3 - (06/05/2026 al 20/05/2026)\n"
            "- Actividad 3\n"
        )
        with open(diligenciar.HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(mock_markdown)

        result = diligenciar.get_all_undiligenced_bitacoras()
        self.assertEqual(result, [1, 3])

    def test_get_all_undiligenced_bitacoras_todas_pendientes(self):
        """Retorna todas las bitácoras si ninguna está diligenciada, en orden."""
        mock_markdown = (
            "# Historico\n\n"
            "## Bitacora numero 1 - (08/04/2026 al 22/04/2026)\n"
            "- Actividad 1\n"
            "## Bitacora numero 2 - (22/04/2026 al 06/05/2026)\n"
            "- Actividad 2\n"
            "## Bitacora numero 3 - (06/05/2026 al 20/05/2026)\n"
            "- Actividad 3\n"
        )
        with open(diligenciar.HISTORICO_PATH, 'w', encoding='utf-8') as f:
            f.write(mock_markdown)

        result = diligenciar.get_all_undiligenced_bitacoras()
        self.assertEqual(result, [1, 2, 3])

    # =========================================================================
    # TESTS UNITARIOS: parse_memory_descriptions validación de errores
    # =========================================================================

    def test_parse_memory_descriptions_archivo_no_existe(self):
        """Falla con FileNotFoundError si el archivo no existe."""
        original = diligenciar.MEMORY_PATH
        try:
            diligenciar.MEMORY_PATH = os.path.join(self.test_dir, "no_existe.md")
            with self.assertRaises(FileNotFoundError):
                diligenciar.parse_memory_descriptions()
        finally:
            diligenciar.MEMORY_PATH = original

    def test_parse_memory_descriptions_sin_bloque_json(self):
        """Falla con ValueError si no hay bloque ```json ... ```."""
        bad_path = os.path.join(self.test_dir, "memory_sin_json.md")
        with open(bad_path, 'w', encoding='utf-8') as f:
            f.write("# Sin bloque JSON\n\nEste archivo no tiene bloque json.")

        original = diligenciar.MEMORY_PATH
        try:
            diligenciar.MEMORY_PATH = bad_path
            with self.assertRaises(ValueError):
                diligenciar.parse_memory_descriptions()
        finally:
            diligenciar.MEMORY_PATH = original

    def test_parse_memory_descriptions_json_malformado(self):
        """Falla si el JSON dentro del bloque está malformado."""
        bad_json_path = os.path.join(self.test_dir, "memory_bad_json.md")
        with open(bad_json_path, 'w', encoding='utf-8') as f:
            f.write("```json\n{ esto no es json valido }\n```")

        original = diligenciar.MEMORY_PATH
        try:
            diligenciar.MEMORY_PATH = bad_json_path
            with self.assertRaises(json.JSONDecodeError):
                diligenciar.parse_memory_descriptions()
        finally:
            diligenciar.MEMORY_PATH = original

    # =========================================================================
    # TESTS DE INTEGRACIÓN EXCEL (actualizados para refactor v2)
    # =========================================================================

    def test_excel_integration(self):
        """
        Prueba de integración actualizada para Refactor v2:
        - Verifica que las actividades están en filas pares (40, 42, 44...)
        - Verifica merge vertical: las filas 40-41 combinadas para actividad 1
        - Verifica que máximo 7 actividades caben (filas 40-53)
        """
        memory_data = diligenciar.parse_memory_descriptions()
        bitacora_data = memory_data["bitacoras"][0]
        exec_date = datetime.date(2026, 5, 23)
        
        # Diligenciar Excel en el directorio de pruebas
        out_path = diligenciar.fill_excel_bitacora(1, bitacora_data, exec_date)
        
        # Verificar que el archivo se creó
        self.assertTrue(os.path.exists(out_path))
        basename = os.path.basename(out_path)
        self.assertTrue(basename.startswith("Bitacora_08-04-2026_22-04-2026"))
        self.assertTrue(basename.endswith(".xlsx"))
        
        # Cargar el archivo generado para inspeccionar sus celdas
        wb = openpyxl.load_workbook(out_path, data_only=True)
        sheet = wb['GFPI-F-147-Formato Bitácora']
        
        # Verificar metadatos rellenados
        self.assertEqual(sheet['E17'].value, 1)
        self.assertEqual(sheet['F17'].value, "DESDE 08/04/2026")
        self.assertEqual(sheet['G17'].value, "HASTA 22/04/2026")
        
        # Verificar que la actividad 1 está en FILA PAR 40
        self.assertEqual(sheet['B40'].value, bitacora_data["actividades"][0]["descripcion"])
        self.assertEqual(sheet['D40'].value, "08/04/2026")
        self.assertEqual(sheet['E40'].value, "10/04/2026")
        self.assertEqual(sheet['F40'].value, bitacora_data["actividades"][0]["evidencia"])
        
        # Verificar MERGE VERTICAL: filas 40-41 combinadas
        merged_ranges = [str(rng) for rng in sheet.merged_cells.ranges]
        self.assertIn('B40:C41', merged_ranges, "B40:C41 debe estar combinado (descripción)")
        self.assertIn('D40:D41', merged_ranges, "D40:D41 debe estar combinado (fecha inicio)")
        self.assertIn('E40:E41', merged_ranges, "E40:E41 debe estar combinado (fecha fin)")
        self.assertIn('F40:F41', merged_ranges, "F40:F41 debe estar combinado (evidencia)")
        
        # Verificar que actividad 2 está en FILA PAR 42
        self.assertEqual(sheet['B42'].value, bitacora_data["actividades"][1]["descripcion"])
        
        # Verificar que actividad 3 está en FILA PAR 44
        self.assertEqual(sheet['B44'].value, bitacora_data["actividades"][2]["descripcion"])
        
        # Verificar que actividad 4 está en FILA PAR 46
        self.assertEqual(sheet['B46'].value, bitacora_data["actividades"][3]["descripcion"])
        
        # Verificar que actividad 5 está en FILA PAR 48
        self.assertEqual(sheet['B48'].value, bitacora_data["actividades"][4]["descripcion"])
        
        # Verificar rango máximo: filas 40 a 53 (7 actividades x 2 filas)
        self.assertEqual(diligenciar.EXCEL_ACT_START_ROW, 40)
        self.assertEqual(diligenciar.EXCEL_ACT_END_ROW, 53)
        self.assertEqual(diligenciar.EXCEL_ACT_ROW_SPAN, 2)
        self.assertEqual(diligenciar.EXCEL_MAX_ACTIVITIES, 7)
        
        wb.close()

    def test_max_7_activities(self):
        """
        Verifica que se trunca a 7 actividades y emite advertencia cuando hay más.
        """
        memory_data = diligenciar.parse_memory_descriptions()
        bitacora_data = memory_data["bitacoras"][0].copy()
        
        # Agregar actividades extra para superar el límite de 7
        extra_activities = bitacora_data["actividades"].copy()
        while len(extra_activities) < 9:
            extra_activities.append(extra_activities[0].copy())
        bitacora_data["actividades"] = extra_activities
        
        exec_date = datetime.date(2026, 5, 23)
        
        # Debe emitir warning por exceder 7 actividades
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            out_path = diligenciar.fill_excel_bitacora(1, bitacora_data, exec_date)
        
        # Verificar que el archivo se creó sin error
        self.assertTrue(os.path.exists(out_path))
        
        # Verificar que solo se escribieron 7 actividades (filas 40-52)
        wb = openpyxl.load_workbook(out_path, data_only=True)
        sheet = wb['GFPI-F-147-Formato Bitácora']
        
        # La actividad 7 debe estar en fila 52 (40 + 6*2)
        self.assertIsNotNone(sheet['B52'].value, "Actividad 7 debe estar en fila 52")
        
        # La fila 54 NO debe tener actividad (fuera del rango)
        self.assertIsNone(sheet['B54'].value, "Fila 54 debe estar vacía (límite excedido)")
        
        wb.close()

    # =========================================================================
    # TESTS DE INTEGRACIÓN WORD MOMENTO 2 (actualizados para refactor v2)
    # =========================================================================

    def test_word_integration_momento_2(self):
        """
        Prueba de integración actualizada para Refactor v2 — Momento 2:
        - P14 (observaciones instructor) NO fue modificado por la función
        - P20 (observaciones co-formador) SÍ fue modificado con retroalimentación
        - P17 (observaciones aprendiz) SÍ tiene contenido nuevo
        - Columna "Compromisos de mejora" tiene texto
        - Texto insertado usa Calibri 9pt
        
        Nota: El template Word puede tener datos de sesiones previas.
        Se verifica que la función NO escribe en P14 pero SÍ en P20 (comportamiento correcto).
        """
        memory_data = diligenciar.parse_memory_descriptions()
        exec_date_str = "08/07/2026"
        
        # Capturar estado ANTES de ejecutar (solo P14, P20 se modifica)
        doc_before = docx.Document(diligenciar.WORD_PATH)
        p14_before = doc_before.paragraphs[14].text
        
        # Diligenciar Momento 2 (retorna la ruta del archivo modificado)
        out_path = diligenciar.process_word_actas(2, memory_data, exec_date_str)
        
        # Cargar el documento MODIFICADO (no el original)
        doc = docx.Document(out_path)
        
        # Verificar marcas X en Tabla 3
        table3 = doc.tables[3]
        self.assertEqual(table3.rows[6].cells[2].text, "X")
        self.assertEqual(table3.rows[20].cells[2].text, "X")
        
        # Verificar fecha de diligenciamiento en Tabla 3
        self.assertEqual(table3.rows[1].cells[8].text, exec_date_str)
        
        # Verificar CAMPOS: P14 NO se modifica, P20 SI se modifica
        self.assertEqual(doc.paragraphs[14].text, p14_before,
                         "P14 (observaciones instructor) NO debe ser modificado en Momento 2")
        expected_coformador = memory_data["actas"]["momento_2"]["observaciones_coformador"]
        self.assertIn(expected_coformador, doc.paragraphs[20].text,
                      "P20 (observaciones co-formador) debe contener retroalimentacion en Momento 2")
        
        # Verificar P17 (observaciones aprendiz) SÍ tiene contenido nuevo
        self.assertIn(memory_data["actas"]["momento_2"]["observaciones_aprendiz"],
                      doc.paragraphs[17].text)
        
        # Verificar observaciones por variable en Tabla 3 (fill_valoraciones_table)
        # Con valoraciones presentes, se escriben las observaciones individuales
        valoraciones = memory_data["actas"]["momento_2"].get("valoraciones", [])
        if valoraciones:
            obs_0 = valoraciones[0].get("observacion", "")
            self.assertIn(obs_0, table3.rows[6].cells[6].text,
                          "Tabla 3 rows[6].cells[6] debe contener la primera observacion por variable")
            obs_4 = valoraciones[4].get("observacion", "")
            self.assertIn(obs_4, table3.rows[17].cells[6].text,
                          "Tabla 3 rows[17].cells[6] debe contener la quinta observacion por variable")
        
        # Verificar Calibri 9pt en P17 (observaciones aprendiz)
        for run in doc.paragraphs[17].runs:
            self.assertEqual(run.font.name, 'Calibri',
                             "P17 debe usar fuente Calibri")
            self.assertEqual(run.font.size, Pt(9),
                             "P17 debe usar tamaño 9pt")
        
        # Verificar Calibri 9pt en marcas X
        for run in table3.rows[6].cells[2].paragraphs[0].runs:
            self.assertEqual(run.font.name, 'Calibri')
            self.assertEqual(run.font.size, Pt(9))
        
        # Verificar párrafo de fecha final P22
        self.assertIn("fecha de diligenciamiento: 08/07/2026", doc.paragraphs[22].text)

    # =========================================================================
    # TESTS DE INTEGRACIÓN WORD MOMENTO 3 (actualizados para refactor v2)
    # =========================================================================

    def test_word_integration_momento_3(self):
        """
        Prueba de integración actualizada para Refactor v2 — Momento 3:
        - Tablas 6 y 7 (co-formador, instructor) SÍ fueron modificadas con retroalimentación
        - Tabla 8 (aprendiz) SÍ tiene contenido nuevo
        - Compromisos de mejora rellenados
        - Calibri 9pt
        
        Nota: El template Word puede tener datos de sesiones previas.
        Se verifica que la función escribe en Tabla6/Tabla7 (comportamiento correcto).
        """
        memory_data = diligenciar.parse_memory_descriptions()
        exec_date_str = "08/10/2026"
        
        # Diligenciar Momento 3 (retorna la ruta del archivo modificado)
        out_path = diligenciar.process_word_actas(3, memory_data, exec_date_str)
        
        # Cargar el documento MODIFICADO (no el original)
        doc = docx.Document(out_path)
        
        # Verificar marcas X en Tabla 5
        table5 = doc.tables[5]
        self.assertEqual(table5.rows[6].cells[3].text, "X")
        self.assertEqual(table5.rows[21].cells[3].text, "X")
        
        # Verificar visitas en Tabla 5
        self.assertEqual(table5.rows[1].cells[10].text, "3")
        
        # Verificar fecha final en Tabla 5
        self.assertEqual(table5.rows[1].cells[7].text, "08/10/2026")
        
        # Verificar que tabla 6 (co-formador) se relleno en R1C1
        expected_coformador_m3 = memory_data["actas"]["momento_3"]["observaciones_coformador"]
        self.assertIn(expected_coformador_m3, doc.tables[6].rows[1].cells[1].text)

        # Verificar que tabla 7 (instructor) se relleno en R1C1
        expected_instructor_m3 = memory_data["actas"]["momento_3"]["observaciones_instructor"]
        self.assertIn(expected_instructor_m3, doc.tables[7].rows[1].cells[1].text)
        
        # Verificar Tabla 8 (aprendiz) SÍ tiene contenido nuevo
        table8 = doc.tables[8]
        self.assertIn(memory_data["actas"]["momento_3"]["observaciones_aprendiz"],
                      table8.rows[2].cells[1].text)
        
        # Verificar compromisos de mejora en Tabla 5
        compromisos = memory_data["actas"]["momento_3"].get("compromisos_mejora", "")
        if compromisos:
            self.assertIn(compromisos, table5.rows[6].cells[6].text)
            self.assertIn(compromisos, table5.rows[17].cells[6].text)
        
        # Verificar Calibri 9pt en Tabla 8 (aprendiz)
        for run in table8.rows[2].cells[1].paragraphs[0].runs:
            self.assertEqual(run.font.name, 'Calibri',
                             "Tabla 8 aprendiz debe usar fuente Calibri")
            self.assertEqual(run.font.size, Pt(9),
                             "Tabla 8 aprendiz debe usar tamaño 9pt")
        
        # Verificar Calibri 9pt en marcas X de Tabla 5
        for run in table5.rows[6].cells[3].paragraphs[0].runs:
            self.assertEqual(run.font.name, 'Calibri')
            self.assertEqual(run.font.size, Pt(9))
        
        # Verificar juicio de evaluación Aprobado en párrafo
        aprobado_paragraph = None
        for p in doc.paragraphs:
            if "Juicio de evaluación" in p.text:
                aprobado_paragraph = p
                break
        self.assertIsNotNone(aprobado_paragraph)
        self.assertIn("Aprobado [X]", aprobado_paragraph.text)

    # =========================================================================
    # NUEVOS TESTS (Refactor v2)
    # =========================================================================

    def test_word_template_not_modified(self):
        """
        Verifica que la plantilla original (actas.docx) NO se modifica
        después de ejecutar process_word_actas. El sistema debe copiar
        la plantilla al directorio de output sin tocar el original.
        """
        memory_data = diligenciar.parse_memory_descriptions()
        exec_date_str = "08/07/2026"
        
        # Capturar hash del archivo plantilla ANTES de ejecutar
        import hashlib
        with open(diligenciar.TEMPLATE_WORD_PATH, 'rb') as f:
            hash_before = hashlib.sha256(f.read()).hexdigest()
        
        # Ejecutar Momento 2
        out_path = diligenciar.process_word_actas(2, memory_data, exec_date_str)
        
        # Verificar que el archivo de salida se creó en el directorio de output
        self.assertTrue(os.path.exists(out_path),
                        "El archivo de salida debe existir")
        self.assertIn("output", out_path,
                      "El archivo de salida debe estar en el directorio de output")
        
        # Verificar que la plantilla original NO fue modificada
        with open(diligenciar.TEMPLATE_WORD_PATH, 'rb') as f:
            hash_after = hashlib.sha256(f.read()).hexdigest()
        
        self.assertEqual(hash_before, hash_after,
                         "La plantilla original NO debe ser modificada tras la ejecución")
        
        # Verificar que el archivo de salida es un documento Word válido
        doc_out = docx.Document(out_path)
        self.assertGreater(len(doc_out.tables), 0,
                           "El archivo de salida debe contener tablas")

    def test_output_folder_creation(self):
        """
        Verifica que se crea la carpeta output/bitacora<N>-<fecha>/ al ejecutar
        fill_excel_bitacora.
        """
        memory_data = diligenciar.parse_memory_descriptions()
        bitacora_data = memory_data["bitacoras"][0]
        exec_date = datetime.date(2026, 5, 23)
        
        # Ejecutar fill_excel_bitacora
        out_path = diligenciar.fill_excel_bitacora(1, bitacora_data, exec_date)
        
        # Verificar que la carpeta de salida existe
        expected_dir = os.path.join(self.test_dir, "output", "bitacora1-2026-05-23")
        self.assertTrue(os.path.isdir(expected_dir),
                        f"La carpeta {expected_dir} debe existir")
        
        # Verificar que el archivo está dentro de la carpeta
        self.assertTrue(out_path.startswith(expected_dir),
                        "El archivo de salida debe estar dentro de la carpeta de output")

    def test_calibri_9_application(self):
        """
        Verifica que la función apply_calibri_9 funciona correctamente
        aplicando Calibri 9pt a un párrafo y a una celda de tabla.
        """
        # Crear un documento Word temporal para pruebas
        test_docx_path = os.path.join(self.test_dir, "test_calibri.docx")
        doc = docx.Document()
        
        # Prueba 1: Aplicar a un párrafo
        p = doc.add_paragraph("Texto de prueba")
        p.add_run(" con Calibri 9")
        
        diligenciar.apply_calibri_9(p)
        
        for run in p.runs:
            self.assertEqual(run.font.name, 'Calibri',
                             "apply_calibri_9 debe establecer fuente Calibri")
            self.assertEqual(run.font.size, Pt(9),
                             "apply_calibri_9 debe establecer tamaño 9pt")
        
        # Prueba 2: Aplicar a una celda de tabla
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "Celda de prueba"
        
        diligenciar.apply_calibri_9(cell)
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                self.assertEqual(run.font.name, 'Calibri',
                                 "apply_calibri_9 en celda debe establecer fuente Calibri")
                self.assertEqual(run.font.size, Pt(9),
                                 "apply_calibri_9 en celda debe establecer tamaño 9pt")
        
        doc.save(test_docx_path)

    def test_word_compromisos_mejora_momento_2(self):
        """
        Verifica que fill_valoraciones_table escribe correctamente en la columna
        'Observaciones / Compromisos de mejora' de la Tabla 3 para Momento 2.
        Con valoraciones presentes, se escriben observaciones individuales (no compromisos_mejora).
        """
        memory_data = diligenciar.parse_memory_descriptions()
        exec_date_str = "08/07/2026"
        
        out_path = diligenciar.process_word_actas(2, memory_data, exec_date_str)
        
        doc = docx.Document(out_path)
        table3 = doc.tables[3]
        
        valoraciones = memory_data["actas"]["momento_2"].get("valoraciones", [])
        if valoraciones:
            # Con valoraciones, rows[6] tiene la primera observacion y rows[17] tiene la quinta
            obs_0 = valoraciones[0].get("observacion", "")
            col6_row6 = table3.rows[6].cells[6].text
            self.assertIn(obs_0, col6_row6,
                          "Tabla 3, fila 6, col 6 debe contener la primera observacion por variable")
            obs_4 = valoraciones[4].get("observacion", "")
            col6_row17 = table3.rows[17].cells[6].text
            self.assertIn(obs_4, col6_row17,
                          "Tabla 3, fila 17, col 6 debe contener la quinta observacion por variable")
        else:
            # Sin valoraciones, debe usar compromisos_mejora como fallback
            compromisos = memory_data["actas"]["momento_2"].get("compromisos_mejora", "")
            if compromisos:
                self.assertIn(compromisos, table3.rows[6].cells[6].text,
                              "Compromisos de mejora deben estar en Tabla 3, fila 6, col 6")
                self.assertIn(compromisos, table3.rows[17].cells[6].text,
                              "Compromisos de mejora deben estar en Tabla 3, fila 17, col 6")

    def test_word_compromisos_mejora_momento_3(self):
        """
        Verifica específicamente que los compromisos de mejora se insertan
        correctamente en la columna 'Observaciones / Compromisos de mejora'
        de la Tabla 5 para Momento 3.
        """
        memory_data = diligenciar.parse_memory_descriptions()
        exec_date_str = "08/10/2026"
        
        out_path = diligenciar.process_word_actas(3, memory_data, exec_date_str)
        
        doc = docx.Document(out_path)
        table5 = doc.tables[5]
        
        compromisos = memory_data["actas"]["momento_3"].get("compromisos_mejora", "")
        if compromisos:
            col6_row6 = table5.rows[6].cells[6].text
            col6_row17 = table5.rows[17].cells[6].text
            self.assertIn(compromisos, col6_row6,
                          "Compromisos de mejora deben estar en Tabla 5, fila 6, col 6")
            self.assertIn(compromisos, col6_row17,
                          "Compromisos de mejora deben estar en Tabla 5, fila 17, col 6")


    # =========================================================================
    # TESTS: Actas state helpers (catch-up logic & state file)
    # =========================================================================

    def test_actas_state_helpers(self):
        """Verifica mark_acta_enviada, reset_actas_state, e inmutabilidad del estado."""
        state = {
            "momento_1": {"enviada": False, "fecha_envio": None},
            "momento_2": {"enviada": False, "fecha_envio": None},
            "momento_3": {"enviada": False, "fecha_envio": None},
        }
        new_state = diligenciar.mark_acta_enviada(state, 2, fecha_iso="2026-06-22")
        # new_state has momento_2 marked
        self.assertTrue(new_state["momento_2"]["enviada"])
        self.assertEqual(new_state["momento_2"]["fecha_envio"], "2026-06-22")
        # Original state is unchanged (immutability)
        self.assertFalse(state["momento_2"]["enviada"])
        self.assertIsNone(state["momento_2"]["fecha_envio"])
        # reset clears everything
        reset = diligenciar.reset_actas_state(new_state)
        for key in ("momento_1", "momento_2", "momento_3"):
            self.assertFalse(reset[key]["enviada"])
            self.assertIsNone(reset[key]["fecha_envio"])

    def test_load_actas_state_missing_file(self):
        """When file doesn't exist, returns fresh state with all enviada=False."""
        missing_path = os.path.join(self.test_dir, "nonexistent_state.json")
        state = diligenciar.load_actas_state(missing_path)
        for key in ("momento_1", "momento_2", "momento_3"):
            self.assertFalse(state[key]["enviada"])
            self.assertIsNone(state[key]["fecha_envio"])

    def test_load_actas_state_invalid_json(self):
        """Malformed JSON returns fresh state and emits warning to stderr."""
        import io
        bad_path = os.path.join(self.test_dir, "bad_state.json")
        with open(bad_path, 'w', encoding='utf-8') as f:
            f.write("{ esto no es json valido }")

        captured_stderr = io.StringIO()
        with patch('sys.stderr', captured_stderr):
            state = diligenciar.load_actas_state(bad_path)

        for key in ("momento_1", "momento_2", "momento_3"):
            self.assertFalse(state[key]["enviada"])
            self.assertIsNone(state[key]["fecha_envio"])
        # Verify warning was emitted
        self.assertIn("Warning", captured_stderr.getvalue())

    def test_load_actas_state_missing_keys(self):
        """JSON with only momento_1 fills missing momento_2/3 with defaults."""
        partial_path = os.path.join(self.test_dir, "partial_state.json")
        with open(partial_path, 'w', encoding='utf-8') as f:
            json.dump({"momento_1": {"enviada": True, "fecha_envio": "2026-04-08"}}, f)

        state = diligenciar.load_actas_state(partial_path)
        self.assertTrue(state["momento_1"]["enviada"])
        self.assertEqual(state["momento_1"]["fecha_envio"], "2026-04-08")
        self.assertFalse(state["momento_2"]["enviada"])
        self.assertIsNone(state["momento_2"]["fecha_envio"])
        self.assertFalse(state["momento_3"]["enviada"])
        self.assertIsNone(state["momento_3"]["fecha_envio"])

    # =========================================================================
    # TESTS: Momento 1 Word integration (new)
    # =========================================================================

    def test_momento_1_word_integration(self):
        """Verifica que process_word_actas(1, ...) rellena Tabla 2 y P10."""
        memory_data = diligenciar.parse_memory_descriptions()
        env_vars = {
            "FECHA_INICIO_ETAPA": "08/04/2026",
            "FECHA_FIN_ETAPA": "22/07/2026",
            "FECHA_AFILIACION_ARL": "01/04/2026",
            "ARL_NUMERO": "123456",
            "HORARIO_ETAPA": "Diurno",
        }
        with patch.dict(os.environ, env_vars):
            out_path = diligenciar.process_word_actas(1, memory_data, "22/06/2026")

        self.assertIsNotNone(out_path)
        doc = docx.Document(out_path)
        table2 = doc.tables[2]

        # R1, C1 — fecha inicio etapa
        self.assertEqual(table2.rows[1].cells[1].text, "08/04/2026")
        # R1, C7 — fecha fin etapa
        self.assertEqual(table2.rows[1].cells[7].text, "22/07/2026")
        # R1, C12 — fecha afiliación ARL
        self.assertEqual(table2.rows[1].cells[12].text, "01/04/2026")
        # R2, C3 — ARL número
        self.assertEqual(table2.rows[2].cells[3].text, "123456")
        # R2, C11 — horario
        self.assertEqual(table2.rows[2].cells[11].text, "Diurno")
        # P10 — fecha de diligenciamiento
        self.assertIn("22/06/2026", doc.paragraphs[10].text)

    def test_momento_1_does_not_touch_firmas(self):
        """Momento 1 no auto-rellena firmas (Tabla 2 R11)."""
        memory_data = diligenciar.parse_memory_descriptions()
        doc_before = docx.Document(diligenciar.WORD_PATH)
        firma_before = [cell.text for cell in doc_before.tables[2].rows[11].cells]

        env_vars = {
            "FECHA_INICIO_ETAPA": "08/04/2026",
            "FECHA_FIN_ETAPA": "22/07/2026",
        }
        with patch.dict(os.environ, env_vars):
            out_path = diligenciar.process_word_actas(1, memory_data, "22/06/2026")

        doc = docx.Document(out_path)
        for idx, cell in enumerate(doc.tables[2].rows[11].cells):
            self.assertEqual(
                cell.text, firma_before[idx],
                f"Tabla 2 R11 cell {idx} (firmas) should not be modified by Momento 1"
            )

    # =========================================================================
    # TESTS: Email module lista de momentos
    # =========================================================================

    def test_email_module_lista_momentos(self):
        """construir_asunto with list of moments produces plural 'Actas Momentos'."""
        subject_12 = email_module.construir_asunto([], [1, 2])
        self.assertIn("Actas Momentos", subject_12)
        self.assertIn("1 y 2", subject_12)

        subject_123 = email_module.construir_asunto([], [1, 2, 3])
        self.assertIn("Actas Momentos", subject_123)
        self.assertIn("1, 2 y 3", subject_123)

    def test_email_module_single_momento_backward_compat(self):
        """construir_asunto with int still works (backward compat)."""
        subject = email_module.construir_asunto([], 2)
        self.assertIn("Acta Momento 2", subject)

    # =========================================================================
    # TESTS: env_validator optional vars
    # =========================================================================

    def test_env_validator_optional_vars(self):
        """Optional vars: absent=ok, invalid=reported, valid=ok."""
        env_path = os.path.join(self.test_dir, ".env_test_optional")
        required_content = (
            "GMAIL_SENDER=test@gmail.com\n"
            "GMAIL_APP_PASSWORD=abcdefghijklmnop\n"
            "EMAIL_DESTINATARIO_PRODUCCION=prod@sena.edu.co\n"
            "EMAIL_DESTINATARIO_PRUEBAS=test@gmail.com\n"
            "EMAIL_CC=cc@gmail.com\n"
            "EMAIL_MODO=pruebas\n"
            "APRENDIZ_NOMBRE=Juan Perez\n"
            "EMPRESA_NOMBRE=Mi Empresa SAS\n"
            "INSTRUCTOR_NOMBRE=Carlos Lopez\n"
        )
        optional_keys = [
            "ACTA_M1_FECHA", "ACTA_M2_FECHA", "ACTA_M3_FECHA",
            "ACTA_VENTANA_DIAS", "FECHA_INICIO_ETAPA", "FECHA_FIN_ETAPA",
            "FECHA_AFILIACION_ARL", "ARL_NUMERO", "HORARIO_ETAPA",
        ]

        # Scenario 1: Required only — should be ok
        with open(env_path, 'w', encoding='utf-8') as f:
            f.write(required_content)
        clear_env = {k: "" for k in optional_keys}
        with patch.dict(os.environ, clear_env):
            ok, missing, _ = env_validator.validate_env(env_path)
        self.assertTrue(ok, "Should be ok with only required vars")
        self.assertEqual(missing, [])

        # Scenario 2: Invalid optional var — should be reported
        with open(env_path, 'w', encoding='utf-8') as f:
            f.write(required_content + "ACTA_M1_FECHA=not-a-date\n")
        with patch.dict(os.environ, clear_env):
            ok, missing, _ = env_validator.validate_env(env_path)
        self.assertFalse(ok, "Should fail with invalid optional var")
        missing_keys = [m["key"] for m in missing]
        self.assertIn("ACTA_M1_FECHA", missing_keys)

        # Scenario 3: Valid optional var — should be ok
        with open(env_path, 'w', encoding='utf-8') as f:
            f.write(required_content + "ACTA_M1_FECHA=2026-04-08\n")
        with patch.dict(os.environ, clear_env):
            ok, missing, _ = env_validator.validate_env(env_path)
        self.assertTrue(ok, "Should be ok with valid optional var")

    # =========================================================================
    # TESTS DE INTEGRACIÓN: Observaciones por variable (valoraciones)
    # =========================================================================

    def test_process_word_momento_2_fills_valoraciones_cells(self):
        """Verifica que M2 escribe las 9 observaciones por variable en Tabla 3 col 6."""
        memory_data = diligenciar.parse_memory_descriptions()
        out_path = diligenciar.process_word_actas(2, memory_data, "22/06/2026")

        doc = docx.Document(out_path)
        table3 = doc.tables[3]

        # Verificar observaciones especificas de las primeras dos valoraciones
        valoraciones = memory_data["actas"]["momento_2"]["valoraciones"]
        self.assertEqual(len(valoraciones), 9, "M2 debe tener 9 valoraciones en el JSON")
        obs_gestion = valoraciones[0]["observacion"]
        obs_creatividad = valoraciones[1]["observacion"]
        self.assertIn(obs_gestion, table3.rows[6].cells[6].text,
                      "Fila 6 col 6 debe contener observacion de 'Gestion de conocimiento'")
        self.assertIn(obs_creatividad, table3.rows[7].cells[6].text,
                      "Fila 7 col 6 debe contener observacion de 'Creatividad y calidad'")

        # Verificar que las 9 celdas (filas 6-9 y 17-21) tienen texto no vacio
        filas_esperadas = [6, 7, 8, 9, 17, 18, 19, 20, 21]
        for fila in filas_esperadas:
            texto = table3.rows[fila].cells[6].text.strip()
            self.assertTrue(len(texto) > 0,
                            f"Tabla 3 rows[{fila}].cells[6] debe tener texto no vacio, obtuvo: '{texto}'")

        # Verificar que NO se incluye "X" en la celda [6] (eso es de la columna 2)
        for fila in filas_esperadas:
            self.assertNotEqual(table3.rows[fila].cells[6].text.strip(), "X",
                                f"Tabla 3 rows[{fila}].cells[6] NO debe ser 'X' (eso pertenece a la columna de marcas)")

    def test_process_word_momento_3_fills_with_fallback(self):
        """Verifica que M3 con valoraciones=[] escribe fallback en rows[6] y rows[17]."""
        memory_data = diligenciar.parse_memory_descriptions()
        out_path = diligenciar.process_word_actas(3, memory_data, "07/10/2026")

        doc = docx.Document(out_path)
        table5 = doc.tables[5]

        # valoraciones en M3 esta vacio, debe usar compromisos_mejora como fallback
        # (que tambien esta vacio en M3, asi que las celdas deben quedar vacias)
        fallback = memory_data["actas"]["momento_3"].get("compromisos_mejora", "")

        # Rows 6 y 17 deben tener el fallback (vacio en este caso)
        col6_row6 = table5.rows[6].cells[6].text.strip()
        col6_row17 = table5.rows[17].cells[6].text.strip()
        self.assertEqual(col6_row6, fallback,
                         "Tabla 5 rows[6].cells[6] debe tener el fallback de M3")
        self.assertEqual(col6_row17, fallback,
                         "Tabla 5 rows[17].cells[6] debe tener el fallback de M3")

        # Las demas celdas (7,8,9,18,19,20,21) deben quedar vacias
        for fila in [7, 8, 9, 18, 19, 20, 21]:
            texto = table5.rows[fila].cells[6].text.strip()
            self.assertEqual(texto, "",
                             f"Tabla 5 rows[{fila}].cells[6] debe estar vacia cuando valoraciones=[]")

    def test_process_word_momento_1_unchanged(self):
        """Regresion: verifica que M1 sigue rellenando Tabla 2 R6-R9 con resultados_aprendizaje."""
        memory_data = diligenciar.parse_memory_descriptions()
        env_vars = {
            "FECHA_INICIO_ETAPA": "08/04/2026",
            "FECHA_FIN_ETAPA": "22/07/2026",
        }
        with patch.dict(os.environ, env_vars):
            out_path = diligenciar.process_word_actas(1, memory_data, "22/06/2026")

        self.assertIsNotNone(out_path)
        doc = docx.Document(out_path)
        table2 = doc.tables[2]

        m1 = memory_data["actas"]["momento_1"]
        self.assertIn(m1["resultados_aprendizaje"], table2.rows[6].cells[2].text,
                      "M1 Tabla 2 R6 debe contener resultados_aprendizaje")
        self.assertIn(m1["actividades_desarrollar"], table2.rows[7].cells[2].text,
                      "M1 Tabla 2 R7 debe contener actividades_desarrollar")
        self.assertIn(m1["evidencias_aprendizaje"], table2.rows[8].cells[2].text,
                      "M1 Tabla 2 R8 debe contener evidencias_aprendizaje")
        self.assertIn(m1["observaciones_adicionales"], table2.rows[9].cells[2].text,
                      "M1 Tabla 2 R9 debe contener observaciones_adicionales")

    # =========================================================================
    # TESTS: Validación del esquema JSON de valoraciones
    # =========================================================================

    def test_memory_has_valoraciones_field(self):
        """Verifica que actas.momento_2.valoraciones exista con 9 entradas validas."""
        memory_data = diligenciar.parse_memory_descriptions()
        self.assertIn("actas", memory_data, "JSON debe tener clave 'actas'")
        self.assertIn("momento_2", memory_data["actas"], "actas debe tener 'momento_2'")
        m2 = memory_data["actas"]["momento_2"]
        self.assertIn("valoraciones", m2, "momento_2 debe tener 'valoraciones'")

        valoraciones = m2["valoraciones"]
        self.assertEqual(len(valoraciones), 9,
                         f"valoraciones debe tener 9 entradas, tiene {len(valoraciones)}")

        for i, val in enumerate(valoraciones):
            self.assertIsInstance(val, dict,
                                 f"valoraciones[{i}] debe ser un dict")
            self.assertIn("variable", val, f"valoraciones[{i}] debe tener 'variable'")
            self.assertTrue(val["variable"].strip(),
                            f"valoraciones[{i}].variable no debe estar vacio")
            self.assertIn("categoria", val, f"valoraciones[{i}] debe tener 'categoria'")
            self.assertTrue(val["categoria"].strip(),
                            f"valoraciones[{i}].categoria no debe estar vacio")
            self.assertIn("observacion", val, f"valoraciones[{i}] debe tener 'observacion'")
            self.assertTrue(val["observacion"].strip(),
                            f"valoraciones[{i}].observacion no debe estar vacio")

    def test_valoraciones_ordered_correctly(self):
        """Verifica que el orden de las valoraciones en M2 sea el del spec."""
        memory_data = diligenciar.parse_memory_descriptions()
        valoraciones = memory_data["actas"]["momento_2"]["valoraciones"]

        expected_order = [
            "Gestion de conocimiento",
            "Creatividad y calidad",
            "Administracion de recursos",
            "Seguridad y salud en el trabajo",
            "Trabajo en equipo",
            "Relaciones interpersonales",
            "Solucion de problemas",
            "Cumplimiento",
            "Organizacion",
        ]

        self.assertEqual(len(valoraciones), len(expected_order),
                         f"Se esperaban {len(expected_order)} valoraciones, se encontraron {len(valoraciones)}")

        for i, expected_var in enumerate(expected_order):
            actual_var = valoraciones[i].get("variable", "")
            self.assertEqual(
                actual_var, expected_var,
                f"valoraciones[{i}] debe ser '{expected_var}', obtuvo '{actual_var}'"
            )


class TestFillValoracionesTable(unittest.TestCase):
    """Tests unitarios para fill_valoraciones_table (observaciones por variable)."""

    def _make_table(self, n_rows=22, n_cols=7):
        """Crea un documento Word temporal con una tabla de n_rows x n_cols."""
        doc = docx.Document()
        table = doc.add_table(rows=n_rows, cols=n_cols)
        return table

    # a) test_fill_valoraciones_writes_all_nine_cells
    def test_fill_valoraciones_writes_all_nine_cells(self):
        """Con 9 valoraciones, las 9 celdas (rows 6,7,8,9,17,18,19,20,21 col 6) deben tener texto."""
        table = self._make_table()
        valoraciones = [
            {"variable": "Gestion de conocimiento", "categoria": "tecnico",
             "observacion": "Demostro apropiacion del conocimiento."},
            {"variable": "Creatividad y calidad", "categoria": "tecnico",
             "observacion": "Aplico buenas practicas de desarrollo."},
            {"variable": "Administracion de recursos", "categoria": "tecnico",
             "observacion": "Fortalecio la gestion de recursos."},
            {"variable": "Seguridad y salud en el trabajo", "categoria": "tecnico",
             "observacion": "Cumplio con los lineamientos de seguridad."},
            {"variable": "Trabajo en equipo", "categoria": "actitudinal",
             "observacion": "Demostro adecuada integracion."},
            {"variable": "Relaciones interpersonales", "categoria": "actitudinal",
             "observacion": "Participo activamente en las reuniones."},
            {"variable": "Solucion de problemas", "categoria": "actitudinal",
             "observacion": "Abordo los desafios con autonomia."},
            {"variable": "Cumplimiento", "categoria": "actitudinal",
             "observacion": "Demostro responsabilidad."},
            {"variable": "Organizacion", "categoria": "actitudinal",
             "observacion": "Evidencio capacidad de organizacion."},
        ]
        fallback = "Compromiso fallback"

        diligenciar.fill_valoraciones_table(table, valoraciones, fallback)

        filas_word = [6, 7, 8, 9, 17, 18, 19, 20, 21]
        for idx, fila in enumerate(filas_word):
            texto = table.rows[fila].cells[6].text
            self.assertEqual(
                texto, valoraciones[idx]["observacion"],
                f"rows[{fila}].cells[6] debe contener la observacion de '{valoraciones[idx]['variable']}', "
                f"obtuvo: '{texto}'"
            )

    # b) test_fill_valoraciones_uses_fallback_when_empty
    def test_fill_valoraciones_uses_fallback_when_empty(self):
        """Con valoraciones=[], fallback debe ir en rows[6] y rows[17] solamente."""
        table = self._make_table()
        fallback = "Compromiso fallback"

        diligenciar.fill_valoraciones_table(table, [], fallback)

        self.assertEqual(table.rows[6].cells[6].text, fallback,
                         "rows[6].cells[6] debe tener fallback cuando valoraciones=[]")
        self.assertEqual(table.rows[17].cells[6].text, fallback,
                         "rows[17].cells[6] debe tener fallback cuando valoraciones=[]")
        # Las demas celdas deben quedar vacias (default del docx)
        for fila in [7, 8, 9, 18, 19, 20, 21]:
            texto = table.rows[fila].cells[6].text.strip()
            self.assertEqual(texto, "",
                             f"rows[{fila}].cells[6] debe estar vacia, obtuvo: '{texto}'")

    # c) test_fill_valoraciones_uses_fallback_when_none
    def test_fill_valoraciones_uses_fallback_when_none(self):
        """Con valoraciones=None, fallback debe ir en rows[6] y rows[17] solamente."""
        table = self._make_table()
        fallback = "Compromiso fallback"

        diligenciar.fill_valoraciones_table(table, None, fallback)

        self.assertEqual(table.rows[6].cells[6].text, fallback,
                         "rows[6].cells[6] debe tener fallback cuando valoraciones=None")
        self.assertEqual(table.rows[17].cells[6].text, fallback,
                         "rows[17].cells[6] debe tener fallback cuando valoraciones=None")
        for fila in [7, 8, 9, 18, 19, 20, 21]:
            texto = table.rows[fila].cells[6].text.strip()
            self.assertEqual(texto, "",
                             f"rows[{fila}].cells[6] debe estar vacia, obtuvo: '{texto}'")

    # d) test_fill_valoraciones_uses_fallback_for_empty_observation
    def test_fill_valoraciones_uses_fallback_for_empty_observation(self):
        """Si la observacion esta vacia/None, debe usar fallback para esa celda."""
        table = self._make_table()
        fallback = "Compromiso fallback"
        valoraciones = [
            {"variable": "Gestion de conocimiento", "categoria": "tecnico",
             "observacion": ""},  # vacia -> fallback
            {"variable": "Creatividad y calidad", "categoria": "tecnico",
             "observacion": "Aplico buenas practicas."},
            {"variable": "Administracion de recursos", "categoria": "tecnico",
             "observacion": None},  # None -> fallback
            {"variable": "Seguridad y salud en el trabajo", "categoria": "tecnico",
             "observacion": "Cumplio con lineamientos."},
            {"variable": "Trabajo en equipo", "categoria": "actitudinal",
             "observacion": "Buena integracion."},
            {"variable": "Relaciones interpersonales", "categoria": "actitudinal",
             "observacion": "Comunicacion asertiva."},
            {"variable": "Solucion de problemas", "categoria": "actitudinal",
             "observacion": "Autonomia analitica."},
            {"variable": "Cumplimiento", "categoria": "actitudinal",
             "observacion": "Responsabilidad demostrada."},
            {"variable": "Organizacion", "categoria": "actitudinal",
             "observacion": "Buena organizacion."},
        ]

        diligenciar.fill_valoraciones_table(table, valoraciones, fallback)

        # rows[6] (idx 0) debe tener fallback (observacion vacia)
        self.assertEqual(table.rows[6].cells[6].text, fallback,
                         "rows[6] debe usar fallback cuando observacion=''")
        # rows[8] (idx 2) debe tener fallback (observacion None)
        self.assertEqual(table.rows[8].cells[6].text, fallback,
                         "rows[8] debe usar fallback cuando observacion=None")
        # rows[7] (idx 1) debe tener su observacion original
        self.assertEqual(table.rows[7].cells[6].text, "Aplico buenas practicas.",
                         "rows[7] debe conservar su observacion original")

    # e) test_fill_valoraciones_handles_dict_without_observacion_key
    def test_fill_valoraciones_handles_dict_without_observacion_key(self):
        """Un dict sin clave 'observacion' no debe explotar; debe usar fallback."""
        table = self._make_table()
        fallback = "Compromiso fallback"
        valoraciones = [
            {"variable": "Gestion de conocimiento", "categoria": "tecnico"},  # sin 'observacion'
            {"variable": "Creatividad y calidad", "categoria": "tecnico",
             "observacion": "Aplico buenas practicas."},
            {"variable": "Administracion de recursos", "categoria": "tecnico",
             "observacion": "Fortalecio gestion."},
            {"variable": "Seguridad y salud en el trabajo", "categoria": "tecnico",
             "observacion": "Cumplio lineamientos."},
            {"variable": "Trabajo en equipo", "categoria": "actitudinal",
             "observacion": "Buena integracion."},
            {"variable": "Relaciones interpersonales", "categoria": "actitudinal",
             "observacion": "Comunicacion asertiva."},
            {"variable": "Solucion de problemas", "categoria": "actitudinal",
             "observacion": "Autonomia analitica."},
            {"variable": "Cumplimiento", "categoria": "actitudinal",
             "observacion": "Responsabilidad."},
            {"variable": "Organizacion", "categoria": "actitudinal",
             "observacion": "Buena organizacion."},
        ]

        # No debe lanzar excepcion
        try:
            diligenciar.fill_valoraciones_table(table, valoraciones, fallback)
        except Exception as e:
            self.fail(f"fill_valoraciones_table no debe explotar con dict sin 'observacion': {e}")

        # rows[6] (idx 0) debe usar fallback
        self.assertEqual(table.rows[6].cells[6].text, fallback,
                         "rows[6] debe usar fallback cuando dict no tiene clave 'observacion'")
        # rows[7] (idx 1) debe tener su observacion
        self.assertEqual(table.rows[7].cells[6].text, "Aplico buenas practicas.",
                         "rows[7] debe conservar su observacion original")

    # f) test_fill_valoraciones_keeps_calibri_style
    def test_fill_valoraciones_keeps_calibri_style(self):
        """Verifica que apply_calibri_9 se aplica a las celdas modificadas."""
        table = self._make_table()
        valoraciones = [
            {"variable": "v1", "categoria": "tecnico", "observacion": "obs1"},
            {"variable": "v2", "categoria": "tecnico", "observacion": "obs2"},
            {"variable": "v3", "categoria": "tecnico", "observacion": "obs3"},
            {"variable": "v4", "categoria": "tecnico", "observacion": "obs4"},
            {"variable": "v5", "categoria": "actitudinal", "observacion": "obs5"},
            {"variable": "v6", "categoria": "actitudinal", "observacion": "obs6"},
            {"variable": "v7", "categoria": "actitudinal", "observacion": "obs7"},
            {"variable": "v8", "categoria": "actitudinal", "observacion": "obs8"},
            {"variable": "v9", "categoria": "actitudinal", "observacion": "obs9"},
        ]

        diligenciar.fill_valoraciones_table(table, valoraciones, "fallback")

        filas_word = [6, 7, 8, 9, 17, 18, 19, 20, 21]
        for fila in filas_word:
            cell = table.rows[fila].cells[6]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self.assertEqual(
                        run.font.name, 'Calibri',
                        f"rows[{fila}].cells[6] debe usar fuente Calibri"
                    )
                    self.assertEqual(
                        run.font.size, Pt(9),
                        f"rows[{fila}].cells[6] debe usar tamano 9pt"
                    )


if __name__ == "__main__":
    unittest.main()
